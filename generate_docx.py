#!/usr/bin/env python3
"""돌고래 헬륨 드론 기술 분석 문서 생성 스크립트"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0, 51, 102)

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0, 76, 153)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)
style_h3.font.color.rgb = RGBColor(0, 102, 153)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_ref(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True


# ═══════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('돌고래 헬륨 드론\n기술 분석 보고서')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Biomimetic Dolphin-Shaped LTA UAV\nfor Indoor Swarm Art Performance')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('2026년 3월\n\n실내 군집 비행 아트 퍼포먼스를 위한\n돌고래 형상 헬륨 드론 설계 및 비행제어 기술 분석')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(60, 60, 60)

doc.add_page_break()

# ═══════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════
doc.add_heading('목차', level=1)
toc_items = [
    '1장. 서론',
    '  1.1 프로젝트 배경 및 목적',
    '  1.2 LTA 비행체 시장 동향',
    '  1.3 생체모방 공중 비행체의 가능성',
    '2장. 헬륨 부력 비행의 기초 원리',
    '  2.1 아르키메데스 원리와 부력 계산',
    '  2.2 헬륨 가스의 물리적 특성',
    '  2.3 부력-중력 평형(Neutral Buoyancy) 설계',
    '  2.4 엔벨로프 체적과 탑재중량 관계',
    '3장. 돌고래 형상의 공기역학적 분석',
    '  3.1 돌고래 유선형 체형의 항력 특성',
    '  3.2 공중에서의 유선형 외피 공력 이점',
    '  3.3 꼬리지느러미(Caudal Fin) 진동 추진 원리',
    '  3.4 등지느러미(Dorsal Fin) — 고정 수동 안정판',
    '  3.5 가슴지느러미(Pectoral Fin) 방향 제어',
    '4장. 기존 제품 및 선행 연구 분석',
    '  4.1 Festo 생체모방 공중 비행체',
    '  4.2 실내 헬륨 비행체 연구 사례',
    '  4.3 군집 드론 아트 퍼포먼스 사례',
    '  4.4 관련 학술 논문 정리',
    '5장. 핵심 기술 요소',
    '  5.1 외피 소재 및 구조',
    '  5.2 구동 메커니즘',
    '  5.3 경량 전자부 설계',
    '  5.4 위치 제어 시스템',
    '  5.5 군집 통신 아키텍처',
    '6장. 비행 제어 시스템 설계',
    '  6.1 제어 시스템 아키텍처',
    '  6.2 자세 제어 (Attitude Control)',
    '  6.3 위치 제어 (Position Control)',
    '  6.4 고도 제어 (Altitude Control)',
    '  6.5 PID 제어기 설계',
    '  6.6 꼬리 진동 추진력 모델링',
    '  6.7 제어 입력 매핑 (Control Allocation)',
    '7장. 군집 비행 제어',
    '  7.1 군집 제어 아키텍처',
    '  7.2 경로 계획 및 안무(Choreography) 설계',
    '  7.3 충돌 회피 알고리즘',
    '  7.4 동기화 및 포메이션 제어',
    '8장. Simulink 모델링 및 제어기 설계',
    '  8.1 Simulink 전체 모델 구조',
    '  8.2 플랜트 모델 (Dolphin Drone Plant)',
    '  8.3 자세 제어기 서브시스템',
    '  8.4 위치 제어기 서브시스템',
    '  8.5 커플링 보상 및 Control Allocation',
    '  8.6 시뮬레이션 시나리오 및 튜닝 절차',
    '9장. Gazebo 연동 시뮬레이션',
    '  9.1 시뮬레이션 환경 구축',
    '  9.2 부력 물리 모델링',
    '  9.3 Simulink-Gazebo-ROS2 연동',
    '  9.4 SITL 기반 비행제어 검증',
    '10장. 설계 사양 및 제작 로드맵',
    '  10.1 목표 사양',
    '  10.2 BOM 초안',
    '  10.3 단계별 개발 로드맵',
    '11장. 결론 및 향후 과제',
    '참고문헌',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    if not item.startswith('  '):
        for r in p.runs:
            r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════
# 1장. 서론
# ═══════════════════════════════════════════
doc.add_heading('1장. 서론', level=1)

doc.add_heading('1.1 프로젝트 배경 및 목적', level=2)
doc.add_paragraph(
    '본 프로젝트는 돌고래의 유선형 형상을 모방한 헬륨 가스 기반 경량 비행체(LTA: Lighter-Than-Air)를 '
    '설계하고, 이를 실내 환경에서 5~10대 규모의 군집 비행을 통해 아트 퍼포먼스를 구현하는 것을 목표로 한다.'
)
doc.add_paragraph(
    '기존의 멀티로터 드론은 소음, 안전성, 배터리 수명 등의 한계로 실내 공연 환경에 적합하지 않다. '
    '반면, 헬륨 부력을 이용한 비행체는 무추력 상태에서도 공중에 체류할 수 있으며, '
    '저소음으로 관객 친화적인 퍼포먼스가 가능하다. 여기에 돌고래의 생체모방(Biomimetic) 설계를 '
    '적용하면 시각적 매력과 유체역학적 효율을 동시에 달성할 수 있다.'
)
add_ref('[참고] Liao, L., & Pasternak, I. (2009). "A review of airship structural research and development." '
        'Progress in Aerospace Sciences, 45(4-5), 83-96.')

doc.add_heading('1.2 LTA 비행체 시장 동향', level=2)
doc.add_paragraph(
    'LTA(Lighter-Than-Air) 비행체 시장은 광고, 감시, 통신 중계, 엔터테인먼트 분야에서 꾸준히 성장하고 있다. '
    '특히 실내 광고용 RC 블림프(Blimp) 시장은 연간 약 15% 성장률을 보이며, '
    '실내 드론쇼 시장은 2025년 기준 약 $2.3B 규모로 추정된다.'
)
doc.add_paragraph(
    '주요 시장 참여자로는 Festo(독일), Skymagic(영국, 구 Verity Studios), '
    'Intel Shooting Star(실외 군집), 그리고 다수의 중국 RC 블림프 제조사가 있다. '
    '이 중 Festo는 생체모방 공중 비행체 분야에서 독보적인 기술력을 보유하고 있으며, '
    'Air Penguin(2009), Air Ray(2007), AirJelly(2008) 등의 프로젝트를 통해 '
    '헬륨 기반 생체모방 비행의 실현 가능성을 입증하였다.'
)
add_ref('[참고] Markets and Markets. (2024). "Unmanned Aerial Vehicle (UAV) Market - Global Forecast."')
add_ref('[참고] Festo AG & Co. KG. (2009). "Air_penguin — Autonomous flight in the air." Bionic Learning Network.')

doc.add_heading('1.3 생체모방 공중 비행체의 가능성', level=2)
doc.add_paragraph(
    '생체모방(Biomimetics)은 자연에서 영감을 얻어 공학적 문제를 해결하는 접근법이다. '
    '수중 생물의 운동 메커니즘을 공중에 적용하는 연구는 Festo의 Bionic Learning Network를 '
    '통해 활발히 진행되어 왔다.'
)
doc.add_paragraph(
    '돌고래는 유체역학적으로 가장 효율적인 체형 중 하나를 가지며, '
    '항력계수(Cd)가 약 0.0036으로 매우 낮다(Fish & Lauder, 2006). '
    '이 유선형 형상을 헬륨 엔벨로프에 적용하면, 공중에서도 낮은 항력으로 에너지 효율적인 '
    '이동이 가능하다. 또한 돌고래의 꼬리지느러미(Caudal fin) 상하 진동 운동은 '
    '경량 서보모터로 구현할 수 있어, 프로펠러 없는 조용한 추진이 가능하다.'
)
add_ref('[참고] Fish, F. E., & Lauder, G. V. (2006). "Passive and active flow control by swimming fishes '
        'and mammals." Annual Review of Fluid Mechanics, 38, 193-224.')
add_ref('[참고] Triantafyllou, M. S., Triantafyllou, G. S., & Yue, D. K. P. (2000). "Hydrodynamics of '
        'fishlike swimming." Annual Review of Fluid Mechanics, 32(1), 33-53.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 2장. 헬륨 부력 비행의 기초 원리
# ═══════════════════════════════════════════
doc.add_heading('2장. 헬륨 부력 비행의 기초 원리', level=1)

doc.add_heading('2.1 아르키메데스 원리와 부력 계산', level=2)
doc.add_paragraph(
    '헬륨 비행체의 부력은 아르키메데스 원리에 기반한다. 비행체가 배제하는 공기의 무게가 '
    '비행체 자체의 무게보다 클 때 양력이 발생한다.'
)
doc.add_paragraph('기본 부력 공식:')
doc.add_paragraph('  F_buoyancy = (ρ_air - ρ_helium) × V × g')
doc.add_paragraph('  여기서:')
doc.add_paragraph('    ρ_air = 1.225 kg/m³ (해수면, 15°C 기준)')
doc.add_paragraph('    ρ_helium = 0.164 kg/m³ (해수면, 15°C 기준)')
doc.add_paragraph('    V = 엔벨로프 체적 (m³)')
doc.add_paragraph('    g = 9.81 m/s²')
doc.add_paragraph()
doc.add_paragraph(
    '따라서 헬륨 1L당 순양력은 약 1.04g이다. 2m급 돌고래 형상 엔벨로프의 '
    '체적을 약 0.3~0.5m³로 가정하면, 순수 부력은 약 310~520g이 된다. '
    '여기서 엔벨로프 자체 무게를 제외한 나머지가 탑재중량(Payload) 여유분이다.'
)
add_ref('[참고] Khoury, G. A. (2012). "Airship Technology." Cambridge University Press, 2nd Edition.')

doc.add_heading('2.2 헬륨 가스의 물리적 특성', level=2)
add_table(
    ['특성', '값', '비고'],
    [
        ['분자량', '4.003 g/mol', '공기(28.97)의 약 1/7'],
        ['밀도 (0°C, 1atm)', '0.1786 kg/m³', ''],
        ['밀도 (20°C, 1atm)', '0.164 kg/m³', '실내 온도 기준'],
        ['비열비 (γ)', '1.66', '단원자 이상기체'],
        ['열전도율', '0.152 W/(m·K)', '공기 대비 약 6배'],
        ['인화성', '없음', '불활성 기체 — 안전'],
        ['순양력 (1L당)', '~1.04 g', '공기 밀도 - 헬륨 밀도'],
        ['확산 속도', '높음', 'Mylar 코팅으로 저감 가능'],
    ]
)
doc.add_paragraph(
    '헬륨은 불활성 기체로 인화성이 없어 실내 공연 환경에서 안전하게 사용할 수 있다. '
    '다만, 헬륨 분자의 크기가 매우 작아(원자 반지름 31pm) 일반 고분자 필름을 통해 '
    '서서히 확산(Permeation)되므로, 엔벨로프 소재의 기밀성이 중요하다.'
)
add_ref('[참고] NIST Chemistry WebBook. "Thermophysical Properties of Helium."')

doc.add_heading('2.3 부력-중력 평형(Neutral Buoyancy) 설계', level=2)
doc.add_paragraph(
    '헬륨 드론의 가장 핵심적인 설계 원칙은 부력-중력 평형(Neutral Buoyancy)이다. '
    '총 부력과 총 중량이 거의 같도록 설계하면, 꼬리지느러미의 미세한 추진력만으로도 '
    '3차원 공간에서의 자유로운 이동이 가능하다.'
)
doc.add_paragraph('설계 평형 조건:')
doc.add_paragraph('  F_buoyancy ≈ W_envelope + W_electronics + W_actuators + W_battery')
doc.add_paragraph()
doc.add_paragraph(
    '실제로는 완벽한 중립 부력보다 약간 양(+)의 부력을 가지도록 설계하는 것이 일반적이다. '
    '이는 전자부품 무게의 오차를 흡수하고, 배터리 방전에 따른 무게 감소를 보상하기 위함이다. '
    '약 5~10%의 양의 부력 여유분을 두며, 이 과잉 부력은 가슴지느러미의 하향 추력으로 상쇄한다.'
)
add_ref('[참고] Mueller, J., et al. (2009). "Development of an aerodynamic model and control law '
        'design for a high altitude airship." AIAA 2009-2810.')

doc.add_heading('2.4 엔벨로프 체적과 탑재중량 관계', level=2)
doc.add_paragraph(
    '2m급 돌고래 형상의 엔벨로프 체적과 예상 탑재중량은 다음과 같다.'
)
add_table(
    ['항목', '값', '산출 근거'],
    [
        ['전체 길이', '약 2.0m', '돌고래 체형 비율 기준'],
        ['최대 직경', '약 0.50m', 'Fineness Ratio 4:1'],
        ['엔벨로프 체적', '약 0.262m³ (262L)', 'Prolate spheroid 공식'],
        ['총 부력', '약 278g', '262L × 1.06g/L'],
        ['엔벨로프 무게', '약 65~80g', 'Mylar 20μm, 표면적 ~2.5m²'],
        ['가용 탑재중량', '약 198~213g', '부력 - 엔벨로프 무게'],
        ['전자부+구동부', '약 30~50g', '경량 설계 시'],
        ['배터리', '약 15~30g', '1S LiPo'],
        ['밸러스트/트림', '약 10~15g', 'CG 조절 + 음성부력 -1g'],
    ]
)
doc.add_paragraph(
    '탑재중량 여유분이 충분하므로, LED 조명이나 추가 센서 탑재도 가능하다. '
    '특히 아트 퍼포먼스용으로 LED 스트립(~10g)을 부착하면 시각적 효과를 극대화할 수 있다.'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 3장. 돌고래 형상의 공기역학적 분석
# ═══════════════════════════════════════════
doc.add_heading('3장. 돌고래 형상의 공기역학적 분석', level=1)

doc.add_heading('3.1 돌고래 유선형 체형의 항력 특성', level=2)
doc.add_paragraph(
    '돌고래(Tursiops truncatus)의 체형은 수백만 년의 자연선택을 통해 최적화된 유선형이다. '
    '체형의 최대 단면적은 전체 길이의 약 30~40% 지점에 위치하며, '
    '이는 층류(Laminar flow) 영역을 최대화하고 난류 천이를 지연시키는 효과가 있다.'
)
doc.add_paragraph(
    'Fish & Hui(1991)의 연구에 따르면, 돌고래의 Fineness Ratio(길이/최대직경)는 약 3.5~5.0이며, '
    '이 범위에서 형상 항력이 최소화된다. NACA 시리즈 에어포일과 유사한 압력 분포를 보이며, '
    '항력계수(Cd)는 약 0.0036~0.0050으로 공학적으로 설계된 유선체와 동등하거나 우수하다.'
)
add_ref('[참고] Fish, F. E., & Hui, C. A. (1991). "Dolphin swimming — a review." '
        'Mammal Review, 21(4), 181-195.')

doc.add_heading('3.2 공중에서의 유선형 외피 공력 이점', level=2)
doc.add_paragraph(
    '헬륨 엔벨로프에 돌고래 형상을 적용할 경우, 단순한 구(Sphere)나 원통형 블림프 대비 '
    '상당한 공력적 이점이 있다.'
)
add_table(
    ['형상', '항력계수 (Cd)', '비고'],
    [
        ['구(Sphere)', '0.47', '기준 형상'],
        ['원통형 블림프', '0.08~0.12', '일반 광고용 블림프'],
        ['유선형(Streamlined)', '0.04~0.06', 'NACA 형상 기반'],
        ['돌고래 형상', '0.035~0.050', 'Fish & Lauder, 2006'],
    ]
)
doc.add_paragraph(
    '실내 환경은 풍속이 낮지만(0~0.5 m/s), 공조 시스템에 의한 미세 기류가 존재한다. '
    '돌고래 유선형은 이러한 외란(Disturbance)에 대한 안정성이 높으며, '
    '전진 비행 시 에너지 효율이 구형 대비 약 5~10배 우수하다.'
)
add_ref('[참고] Hoerner, S. F. (1965). "Fluid-Dynamic Drag." Published by the author.')

doc.add_heading('3.3 꼬리지느러미(Caudal Fin) 진동 추진 원리', level=2)
doc.add_paragraph(
    '돌고래는 꼬리지느러미(Caudal fin)의 상하 진동(Dorso-ventral oscillation)을 통해 추진력을 얻는다. '
    '이는 어류의 좌우 진동(Lateral undulation)과 구별되는 해양 포유류 고유의 운동 방식이다.'
)
doc.add_paragraph('추진력 모델:')
doc.add_paragraph('  F_thrust = 0.5 × ρ × A_fin × V²_tip × C_T')
doc.add_paragraph('  여기서:')
doc.add_paragraph('    A_fin = 꼬리지느러미 면적')
doc.add_paragraph('    V_tip = 지느러미 끝단 속도 = 2πfA (f: 주파수, A: 진폭)')
doc.add_paragraph('    C_T = 추력계수 (Strouhal Number에 의존)')
doc.add_paragraph()
doc.add_paragraph(
    'Strouhal Number(St = fA/U)가 0.2~0.4 범위일 때 추진 효율이 최대가 되며, '
    '이는 자연계의 유영 생물에서 보편적으로 관찰되는 값이다(Taylor et al., 2003). '
    '헬륨 드론에서는 저속(U < 1 m/s) 환경이므로, 낮은 주파수(1~3 Hz)와 '
    '큰 진폭(±30°)의 조합이 적절하다.'
)
add_ref('[참고] Taylor, G. K., Nudds, R. L., & Thomas, A. L. R. (2003). "Flying and swimming animals '
        'cruise at a Strouhal number tuned for high power efficiency." Nature, 425(6959), 707-711.')
add_ref('[참고] Sfakiotakis, M., Lane, D. M., & Davies, J. B. C. (1999). "Review of fish swimming modes '
        'for aquatic locomotion." IEEE Journal of Oceanic Engineering, 24(2), 237-252.')

doc.add_heading('3.4 등지느러미(Dorsal Fin) — 고정 수동 안정판', level=2)
doc.add_paragraph(
    '돌고래의 등지느러미(Dorsal fin)는 추진 기관이 아니라 안정화 장치이다. '
    '생물학적으로는 Roll 안정성(용골 효과)과 Yaw 안정성(직진성 유지)을 담당하며, '
    '비행기의 수직꼬리날개(Vertical stabilizer)와 동일한 원리로 작동한다.'
)
doc.add_paragraph()
doc.add_paragraph(
    '본 프로젝트에서는 등지느러미를 고정 수동 안정판으로 설계한다. '
    '능동 러더(Active Rudder) 방식도 검토하였으나, 헬륨 드론의 저속 특성(0.3~0.5 m/s)에서는 '
    '러더에 충분한 공기역학적 조향력이 발생하지 않아(동압 q ≈ 0.05 Pa로 극히 낮음) '
    '실질적인 Yaw 제어가 불가능하다. 따라서 Yaw 제어는 꼬리지느러미의 좌우 편향으로 수행하고, '
    '등지느러미는 수동적 안정성 확보에 집중한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 고정 등지느러미의 역할')
doc.add_paragraph('  • Yaw 안정성: 직진 비행 시 풍향각(Sideslip) 발생 시 복원 모멘트 제공')
doc.add_paragraph('  • Roll 안정성: 용골(Keel) 효과로 좌우 흔들림 감쇠')
doc.add_paragraph('  • 꼬리 진동 시 발생하는 Yaw 진동을 수동적으로 감쇠')
doc.add_paragraph('  • 서보 불필요 → 무게/전력/복잡도 추가 없음')
doc.add_paragraph()
doc.add_paragraph('■ 능동 러더를 채택하지 않은 이유')
doc.add_paragraph('  러더 조향력: F_rudder = 0.5 × ρ × V² × S × C_L × sin(δ)')
doc.add_paragraph('  V = 0.3 m/s, S = 0.015 m² 일 때: F ≈ 0.04 gf → 사실상 조향력 없음')
doc.add_paragraph('  반면 꼬리 좌우 편향은 추력 기반이므로 속도와 무관하게 Yaw 제어 가능')
doc.add_paragraph()
doc.add_paragraph('■ 등지느러미 구조')
doc.add_paragraph('  • 소재: 엔벨로프와 일체형 Mylar + 내부 카본 로드 뼈대')
doc.add_paragraph('  • 높이: 약 10~15cm')
doc.add_paragraph('  • 무게: ~1~2g (서보 없음)')
doc.add_paragraph('  • 고정 방식: 엔벨로프 열접합 일체형')
add_ref('[참고] Fish, F. E., & Rohr, J. J. (1999). "Review of dolphin hydrodynamics and swimming performance." '
        'SPAWAR Systems Center Technical Report 1801.')

doc.add_heading('3.5 가슴지느러미(Pectoral Fin) 방향 제어', level=2)
doc.add_paragraph(
    '돌고래의 가슴지느러미(Pectoral fin)는 추진보다는 조향과 자세 제어에 사용된다. '
    '양쪽 가슴지느러미의 각도를 독립적으로 제어함으로써:'
)
doc.add_paragraph('  • 좌우 비대칭 → Roll 제어')
doc.add_paragraph('  • 동시 상향/하향 → Pitch 및 고도 제어')
doc.add_paragraph('  • 비대칭 항력 → 보조적 Yaw 제어')
doc.add_paragraph()
doc.add_paragraph(
    '헬륨 드론에서는 각 가슴지느러미를 경량 서보모터(~3g)로 구동하며, '
    '±45° 범위의 각도 제어가 가능하도록 설계한다. '
    '부력과 거의 평형 상태이므로, 작은 제어력으로도 충분한 기동이 가능하다.'
)
add_ref('[참고] Fish, F. E. (2004). "Structure and mechanics of nonpiscine control surfaces." '
        'IEEE Journal of Oceanic Engineering, 29(3), 605-621.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 4장. 기존 제품 및 선행 연구 분석
# ═══════════════════════════════════════════
doc.add_heading('4장. 기존 제품 및 선행 연구 분석', level=1)

doc.add_heading('4.1 Festo 생체모방 공중 비행체', level=2)
doc.add_paragraph(
    'Festo AG의 Bionic Learning Network는 생체모방 공중 비행체 분야에서 가장 앞선 사례를 보여준다.'
)
add_table(
    ['제품명', '연도', '형태', '핵심 기술', '크기/무게'],
    [
        ['Air Ray', '2007', '가오리', '헬륨 부력 + 날개 플래핑', '길이 4.2m, 1.6kg'],
        ['AirJelly', '2008', '해파리', '헬륨 부력 + 텐타클 추진', '직경 1.35m, 1.3kg'],
        ['Air Penguin', '2009', '펭귄', '헬륨 부력 + 핀 추진 + 자율비행', '길이 1.0m, ~1kg'],
        ['eMotionSpheres', '2015', '구체', '헬륨 + 프로펠러 + 군집', '직경 0.5m, 0.3kg'],
        ['FreeMotionHandling', '2017', '구체', '헬륨 + 그리퍼 + 실내GPS', '직경 0.7m, ~1kg'],
        ['BionicFinWave', '2018', '물고기', '수중, 핀 물결 추진', '길이 0.37m, ~0.5kg'],
    ]
)
doc.add_paragraph(
    '특히 Air Penguin은 본 프로젝트와 가장 유사한 사례로, 3D 비행이 가능한 자율 헬륨 비행체이다. '
    '두 개의 핀(가슴지느러미)을 이용하여 추진과 조향을 동시에 수행하며, '
    '머리 부분의 방향타(Rudder)로 추가적인 방향 제어를 한다. '
    'UWB(Ultra-Wideband) 기반의 실내 위치추적 시스템을 탑재하여 자율 비행을 구현하였다.'
)
add_ref('[참고] Festo AG & Co. KG. (2007-2018). Bionic Learning Network Technical Reports. '
        'https://www.festo.com/group/en/cms/10156.htm')
add_ref('[참고] Festo AG. (2009). "Air_penguin — Autonomous flying and gliding." Technical Documentation.')

doc.add_heading('4.2 실내 헬륨 비행체 연구 사례', level=2)
doc.add_paragraph(
    '학술 연구 분야에서도 실내 헬륨 비행체에 대한 연구가 활발히 진행되어 왔다.'
)
doc.add_paragraph(
    '• Georgia Tech — GTMARS(Georgia Tech Miniature Autonomous Robotic Airship): '
    '실내 자율 비행을 위한 소형 블림프 플랫폼. '
    'Rao et al.(2003)은 초음파 센서 기반의 장애물 회피 시스템을 개발하였다.'
)
doc.add_paragraph(
    '• MIT — 실내 블림프 제어 연구: Van der Zwaan et al.(2000)은 '
    '비전 기반 실내 블림프 자율 비행 시스템을 제안하였으며, '
    '단안 카메라를 이용한 위치 추정을 구현하였다.'
)
doc.add_paragraph(
    '• ETH Zurich — 경량 실내 비행체 연구: Zufferey et al.(2006)의 '
    '경량 자율 비행체 연구는 10g급 전자부 설계의 가능성을 보여주었다.'
)
add_ref('[참고] Rao, J., et al. (2003). "GTMARS: A miniature autonomous robotic airship." '
        'IEEE/RSJ International Conference on Intelligent Robots and Systems.')
add_ref('[참고] Van der Zwaan, S., et al. (2000). "Vision based station keeping and docking '
        'for an aerial blimp." IEEE/RSJ IROS 2000.')

doc.add_heading('4.3 군집 드론 아트 퍼포먼스 사례', level=2)
doc.add_paragraph(
    '군집 드론 아트 퍼포먼스는 최근 급격히 성장하고 있는 분야이다.'
)
add_table(
    ['사례', '드론 수', '환경', '위치추적', '비고'],
    [
        ['Intel Shooting Star', '최대 2,000+', '실외', 'RTK-GPS', '올림픽/월드컵 공연'],
        ['Verity Studios (Skymagic)', '50~200', '실내', '외부 카메라 트래킹', 'Cirque du Soleil 협업'],
        ['Festo eMotionSpheres', '5~10', '실내', 'UWB', 'Hannover Messe 시연'],
        ['Collmot Robotics', '최대 500', '실외', 'RTK-GPS', '유럽 드론쇼'],
        ['ETH Zurich Crazyswarm', '49', '실내', 'Vicon 모캡', '연구/시연용'],
    ]
)
doc.add_paragraph(
    '실내 군집 비행에서는 Verity Studios(현 Skymagic)의 사례가 가장 참고할 만하다. '
    'Cirque du Soleil과의 협업을 통해 수십 대의 소형 드론이 공연장에서 안무를 수행하였으며, '
    '외부 카메라 기반 위치추적과 중앙 집중형 제어 아키텍처를 채택하였다.'
)
add_ref('[참고] Preiss, J. A., et al. (2017). "Crazyswarm: A large nano-quadcopter swarm." '
        'IEEE International Conference on Robotics and Automation (ICRA).')
add_ref('[참고] Augugliaro, F., et al. (2014). "The Flight Assembled Architecture installation." '
        'IEEE Robotics & Automation Magazine, 21(3), 34-44.')

doc.add_heading('4.4 관련 학술 논문 정리', level=2)
doc.add_paragraph('본 프로젝트와 직접적으로 관련된 핵심 논문을 분류별로 정리한다.')
add_table(
    ['분류', '논문/저자', '연도', '핵심 내용'],
    [
        ['LTA 비행체 역학', 'Li, Y., et al. "Dynamic modeling of an airship"', '2011',
         '6-DOF 비행체 운동방정식, 부력항 포함 모델링'],
        ['핀 추진 공력', 'Bandyopadhyay, P. R. "Maneuvering hydrodynamics of fish and small UAVs"', '2005',
         '핀 추진 유체역학, Strouhal Number 최적화'],
        ['생체모방 비행', 'Shyy, W., et al. "Flapping and flexible wings for biological and micro air vehicles"', '2010',
         '생체모방 날개 설계, 유연 구조 역학'],
        ['실내 블림프 제어', 'Fedorenko, R., & Krukhmalev, V. "Indoor autonomous airship control"', '2016',
         '실내 블림프 자세/위치 제어 알고리즘'],
        ['군집 제어', 'Vásárhelyi, G., et al. "Optimized flocking of autonomous drones"', '2018',
         '대규모 군집 비행 알고리즘, 자기조직화'],
        ['경로 계획', 'Mellinger, D., & Kumar, V. "Minimum snap trajectory generation"', '2011',
         '최소 스냅 궤적 생성, 다항식 궤적 최적화'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 5장. 핵심 기술 요소
# ═══════════════════════════════════════════
doc.add_heading('5장. 핵심 기술 요소', level=1)

doc.add_heading('5.1 외피 소재 및 구조', level=2)
doc.add_paragraph(
    '엔벨로프(외피)는 헬륨 가스를 담는 기밀 구조체로, 비행체의 형상과 부력을 결정하는 핵심 요소이다.'
)
add_table(
    ['소재', '두께', '면밀도', '기밀성', '가공성', '비고'],
    [
        ['Mylar (PET)', '12~25μm', '15~30 g/m²', '★★★★★', '★★★', '가장 일반적, 열접합'],
        ['Nylon PA6', '15~30μm', '17~35 g/m²', '★★★★', '★★★★', '유연성 우수'],
        ['TPU Film', '20~50μm', '24~60 g/m²', '★★★★', '★★★★★', '탄성, 반복사용 적합'],
        ['EVOH 코팅 Nylon', '15~25μm', '18~30 g/m²', '★★★★★', '★★★', '최고 기밀성'],
        ['Al-Mylar 복합', '15~20μm', '20~28 g/m²', '★★★★★', '★★★', 'Al 코팅으로 확산 저감'],
    ]
)
doc.add_paragraph(
    '본 프로젝트에서는 Mylar(PET) 12~20μm 필름을 기본 소재로 권장한다. '
    '가격이 저렴하고, 열접합(Heat sealing)이 가능하며, 헬륨 투과율이 낮아 '
    '수일 간 부력을 유지할 수 있다. 돌고래 형상은 3D 패턴을 전개도로 설계한 후 '
    '열접합으로 조립하는 방식이 적합하다.'
)
add_ref('[참고] Colozza, A. (2003). "Initial feasibility assessment of a high altitude long endurance '
        'airship." NASA/CR-2003-212724.')

doc.add_heading('5.2 구동 메커니즘', level=2)
doc.add_paragraph(
    '돌고래 헬륨 드론의 능동 구동부는 꼬리지느러미(서보 2개)와 가슴지느러미(서보 2개)로 구성된다. '
    '등지느러미는 고정 수동 안정판으로, 별도의 액추에이터 없이 Yaw/Roll 안정성을 제공한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 꼬리지느러미 구동 (추진 + Yaw 조향)')
doc.add_paragraph('  • 서보모터 1개: 상하 진동 운동 (추진력 생성)')
doc.add_paragraph('  • 서보모터 1개: 좌우 편향 (Yaw 방향 조향)')
doc.add_paragraph('  • 진동 주파수: 1~3 Hz')
doc.add_paragraph('  • 진동 진폭: ±20°~±40°')
doc.add_paragraph('  • 추진력 제어: 주파수 × 진폭으로 조절')
doc.add_paragraph('  • Yaw 제어: 좌우 편향 각도로 조절 (추력 기반, 저속에서도 유효)')
doc.add_paragraph()
doc.add_paragraph('■ 등지느러미 (고정 — 수동 안정판)')
doc.add_paragraph('  • 서보 없음 — 고정 구조물')
doc.add_paragraph('  • Yaw/Roll 방향 수동 감쇠 효과')
doc.add_paragraph('  • 무게: ~1~2g (카본 로드 뼈대 + Mylar)')
doc.add_paragraph()
doc.add_paragraph('■ 가슴지느러미 구동 (Roll/Pitch/고도)')
doc.add_paragraph('  • 서보모터 2개 (좌/우 독립 제어)')
doc.add_paragraph('  • 각도 범위: ±45°')
doc.add_paragraph('  • Roll 제어: 좌우 비대칭 각도')
doc.add_paragraph('  • Pitch/고도 제어: 동시 상향/하향')
doc.add_paragraph()
doc.add_paragraph('■ 구동부 구성도')
doc.add_paragraph()
doc.add_paragraph('           등지느러미 (고정, 수동 안정판)')
doc.add_paragraph('                │')
doc.add_paragraph('           ┌────△────┐')
doc.add_paragraph('          /     │     \\')
doc.add_paragraph('     ◁━━○      │      ○━━▷   ← 가슴지느러미 L/R (서보 2개)')
doc.add_paragraph('          \\     │     /')
doc.add_paragraph('           └────┼────┘')
doc.add_paragraph('                │')
doc.add_paragraph('                ◇━━━━━━◇      ← 꼬리 서보 2개 (상하:추진, 좌우:Yaw)')
doc.add_paragraph()
add_table(
    ['구동부', '액추에이터', '무게', '제어 변수', '역할'],
    [
        ['꼬리 (상하)', '서보 SG51R 또는 동급', '~3.5g', '주파수, 진폭', '추진'],
        ['꼬리 (좌우)', '서보 SG51R 또는 동급', '~3.5g', '편향 각도', 'Yaw 조향'],
        ['좌측 가슴지느러미', '서보 SG51R 또는 동급', '~3.5g', '각도 (±45°)', 'Roll/Pitch/고도'],
        ['우측 가슴지느러미', '서보 SG51R 또는 동급', '~3.5g', '각도 (±45°)', 'Roll/Pitch/고도'],
        ['등지느러미', '없음 (고정)', '~1~2g', '-', '수동 안정판'],
        ['합계', '서보 4개', '~14g(+2g 등지느러미)', '', ''],
    ]
)
doc.add_paragraph(
    '꼬리가 추진과 Yaw를 동시에 담당하므로 추진-조향 간 커플링이 존재하나, '
    '이는 제어기 설계(Simulink)에서 크로스 커플링 보상항으로 해결한다. '
    '저속 헬륨 드론에서는 꼬리 편향이 추력 기반(속도 무관)이므로 '
    '정지/저속 상태에서도 Yaw 제어가 가능하다는 장점이 능동 러더 방식 대비 크다.'
)

doc.add_heading('5.3 경량 전자부 설계', level=2)
doc.add_paragraph(
    '군집 비행을 위한 경량 전자부 구성은 다음과 같다.'
)
add_table(
    ['부품', '모델', '무게', '역할', '전력'],
    [
        ['MCU', 'ESP32-C3 Mini', '~2g', 'WiFi 통신 + 자세제어', '~80mA'],
        ['IMU', 'MPU6050 / ICM-20948', '~1g', '가속도/자이로/지자기', '~3.5mA'],
        ['서보 드라이버', 'PCA9685 (선택)', '~1g', '4채널 PWM 출력', '~10mA'],
        ['배터리', '1S LiPo 300mAh', '~8g', '전원 공급', '-'],
        ['배터리', '1S LiPo 600mAh', '~16g', '전원 공급 (장시간)', '-'],
        ['전압 레귤레이터', 'AMS1117-3.3', '<1g', '3.3V 안정화', '-'],
        ['IR 마커', '반사구 3~4개', '<1g', '외부 트래킹용', '-'],
        ['LED (선택)', 'WS2812B 5개', '~3g', '아트 퍼포먼스 조명', '~100mA'],
        ['합계 (기본)', '-', '~13g', '-', '~94mA'],
        ['합계 (LED 포함)', '-', '~16g', '-', '~194mA'],
    ]
)
doc.add_paragraph(
    '배터리 300mAh 기준 약 3시간, 600mAh 기준 약 6시간의 비행이 가능하다 (기본 구성 기준). '
    'LED 사용 시에는 약 1.5~3시간으로 줄어든다.'
)

doc.add_heading('5.4 위치 제어 시스템 — 외부 카메라 트래킹', level=2)
doc.add_paragraph(
    '군집 아트 퍼포먼스를 위한 위치 제어는 외부 카메라 트래킹 시스템을 채택한다. '
    '이 방식은 기체에 무게 부담이 거의 없으며, 다수의 기체를 동시에 추적할 수 있다.'
)
doc.add_paragraph('■ Option A: 상용 모션캡처 시스템')
doc.add_paragraph('  • OptiTrack / Vicon 시스템')
doc.add_paragraph('  • 카메라 6~8대로 비행 공간 커버')
doc.add_paragraph('  • 정밀도: ±1mm 이하')
doc.add_paragraph('  • 동시 추적: 수십 개 객체')
doc.add_paragraph('  • 비용: $15,000~$30,000+')
doc.add_paragraph()
doc.add_paragraph('■ Option B: 저가 DIY 카메라 트래킹')
doc.add_paragraph('  • 산업용 카메라 4~6대 + OpenCV')
doc.add_paragraph('  • ArUco 마커 또는 IR LED 마커')
doc.add_paragraph('  • 정밀도: ±2~5cm (아트 퍼포먼스에 충분)')
doc.add_paragraph('  • 비용: $500~$2,000')
doc.add_paragraph()
doc.add_paragraph(
    '두 방식 모두 기체에는 반사 마커(~1g 미만)만 부착하면 되므로, '
    '탑재중량에 미치는 영향이 거의 없다. 이는 UWB 태그(~4g)나 GPS 모듈(~10g) 대비 '
    '큰 장점이다.'
)
add_ref('[참고] Krajník, T., et al. (2014). "A practical multirobot localization system." '
        'Journal of Intelligent & Robotic Systems, 76(3), 539-562.')

doc.add_heading('5.5 군집 통신 아키텍처', level=2)
doc.add_paragraph(
    '5~10대 규모의 군집 비행을 위한 통신 아키텍처는 중앙 집중형(Centralized) 구조를 채택한다.'
)
doc.add_paragraph('■ 중앙 집중형 아키텍처 (Star Topology)')
doc.add_paragraph('  • Ground Station(PC)이 모든 드론의 위치를 파악하고 명령을 전송')
doc.add_paragraph('  • 통신: WiFi (ESP32 내장)')
doc.add_paragraph('  • 프로토콜: UDP (저지연, 경량)')
doc.add_paragraph('  • 갱신 주기: 10~20 Hz')
doc.add_paragraph('  • 데이터 패킷: 목표 위치(x,y,z) + 자세(roll,pitch,yaw) = ~24 bytes')
doc.add_paragraph()
doc.add_paragraph(
    '10대 × 20Hz × 24bytes = 약 4.8 KB/s로, WiFi 대역폭 대비 매우 경량이다. '
    '지연시간(Latency)은 WiFi 환경에서 약 5~20ms로, 저속 헬륨 드론 제어에 충분하다.'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 6장. 비행 제어 시스템 설계 (핵심 장)
# ═══════════════════════════════════════════
doc.add_heading('6장. 비행 제어 시스템 설계', level=1)

doc.add_paragraph(
    '본 장에서는 돌고래 헬륨 드론의 비행제어 시스템을 상세히 다룬다. '
    '기존 멀티로터 드론의 PX4/Ardupilot과 같은 상용 비행제어기는 '
    '모터 RPM 기반 제어를 전제로 설계되어 있어, '
    '핀(지느러미) 추진 기반의 헬륨 드론에는 적합하지 않다. '
    '따라서 커스텀 비행제어기를 ESP32 기반으로 설계한다.'
)

doc.add_heading('6.1 제어 시스템 아키텍처', level=2)
doc.add_paragraph(
    '비행제어 시스템은 2단 계층 구조(Cascaded Control)로 설계한다.'
)
doc.add_paragraph('■ 외부 루프 (Outer Loop) — 위치 제어기 [Ground Station PC]')
doc.add_paragraph('  • 실행 주기: 10~20 Hz')
doc.add_paragraph('  • 입력: 현재 위치 (모션캡처), 목표 위치 (안무 데이터)')
doc.add_paragraph('  • 출력: 목표 자세 (Roll, Pitch, Yaw) + 목표 추진력')
doc.add_paragraph('  • 알고리즘: PID 또는 MPC (Model Predictive Control)')
doc.add_paragraph()
doc.add_paragraph('■ 내부 루프 (Inner Loop) — 자세 제어기 [ESP32 온보드]')
doc.add_paragraph('  • 실행 주기: 50~100 Hz')
doc.add_paragraph('  • 입력: 현재 자세 (IMU), 목표 자세 (외부 루프에서 수신)')
doc.add_paragraph('  • 출력: 서보 명령 (꼬리 주파수/진폭, 지느러미 각도)')
doc.add_paragraph('  • 알고리즘: PID')
doc.add_paragraph()
doc.add_paragraph('제어 흐름도:')
doc.add_paragraph()
doc.add_paragraph('  [안무 데이터] → [경로 계획기] → [위치 제어기(PC)]')
doc.add_paragraph('       ↑                              ↓ WiFi (목표 자세)')
doc.add_paragraph('  [모션캡처] ──────────→ [자세 제어기(ESP32)]')
doc.add_paragraph('       (위치 피드백)              ↓')
doc.add_paragraph('                          [서보 출력]')
doc.add_paragraph('                     ┌──────┼──────┐')
doc.add_paragraph('                   꼬리(상하) 꼬리(좌우) 가슴지느러미(L/R)')
doc.add_paragraph('                    [추진]    [Yaw]    [Roll/Pitch/고도]')
doc.add_paragraph()
doc.add_paragraph(
    '이 2단 구조의 핵심 이점은 역할 분리이다. '
    'PC는 연산량이 많은 위치 제어와 경로 계획을 담당하고, '
    'ESP32는 빠른 주기의 자세 안정화만 담당하여 경량 MCU로도 충분한 성능을 발휘한다.'
)
add_ref('[참고] Beard, R. W., & McLain, T. W. (2012). "Small Unmanned Aircraft: Theory and Practice." '
        'Princeton University Press.')

doc.add_heading('6.2 자세 제어 (Attitude Control)', level=2)
doc.add_paragraph(
    '자세 제어는 비행 안정성의 근간이다. IMU(MPU6050)로부터 Roll, Pitch, Yaw 각도를 측정하고, '
    '목표 자세와의 오차를 최소화하도록 서보를 구동한다.'
)
doc.add_paragraph('■ 자세 추정 (Attitude Estimation)')
doc.add_paragraph(
    '  MPU6050의 가속도계와 자이로스코프 데이터를 상보 필터(Complementary Filter) 또는 '
    '  Madgwick 필터로 융합하여 자세를 추정한다.'
)
doc.add_paragraph()
doc.add_paragraph('  Madgwick 필터 장점:')
doc.add_paragraph('    - 연산량이 적어 ESP32에서 실시간 구동 가능')
doc.add_paragraph('    - 자이로 드리프트 보정 우수')
doc.add_paragraph('    - 단일 파라미터(β) 튜닝으로 간편')
doc.add_paragraph()
doc.add_paragraph('■ Roll 제어')
doc.add_paragraph('  제어 입력: 좌/우 가슴지느러미의 차동 각도 (δ_L - δ_R)')
doc.add_paragraph('  Roll 오차 → PID → Δδ_pectoral')
doc.add_paragraph('    δ_L = δ_base + Δδ_pectoral')
doc.add_paragraph('    δ_R = δ_base - Δδ_pectoral')
doc.add_paragraph()
doc.add_paragraph('■ Pitch 제어')
doc.add_paragraph('  제어 입력: 좌/우 가슴지느러미의 공통 각도 (δ_L + δ_R)/2')
doc.add_paragraph('  Pitch 오차 → PID → δ_common')
doc.add_paragraph('  양쪽 지느러미를 동시에 같은 방향으로 틸트하여 피치 모멘트 발생')
doc.add_paragraph()
doc.add_paragraph('■ Yaw 제어')
doc.add_paragraph('  제어 입력: 꼬리의 좌우 편향 각도 (δ_tail_yaw)')
doc.add_paragraph('  Yaw 오차 → PID → δ_tail_yaw')
doc.add_paragraph('  꼬리 진동의 좌우 비대칭(bias)으로 Yaw 모멘트 생성')
doc.add_paragraph('  추력 기반이므로 저속/정지 상태에서도 Yaw 제어 가능')
doc.add_paragraph('  등지느러미(고정)가 수동적으로 Yaw 진동을 감쇠하여 안정성 보조')
doc.add_paragraph()
doc.add_paragraph('■ 추진-조향 커플링 보상')
doc.add_paragraph('  꼬리가 추진(상하)과 Yaw(좌우)를 동시에 수행하므로 커플링이 존재한다.')
doc.add_paragraph('  보상 전략:')
doc.add_paragraph('    1. Yaw 편향 시 추진력 감소분을 주파수/진폭 증가로 보상')
doc.add_paragraph('    2. Simulink에서 크로스 커플링 전달함수를 모델링하여 피드포워드 보상')
doc.add_paragraph('    3. 급격한 Yaw 변경 시 추진력을 일시적으로 증가시키는 스케줄링')
doc.add_paragraph()
add_ref('[참고] Madgwick, S. O. H., Harrison, A. J. L., & Vaidyanathan, R. (2011). '
        '"Estimation of IMU and MARG orientation using a gradient descent algorithm." '
        'IEEE International Conference on Rehabilitation Robotics.')

doc.add_heading('6.3 위치 제어 (Position Control)', level=2)
doc.add_paragraph(
    '위치 제어기는 Ground Station PC에서 실행되며, 모션캡처 시스템으로부터 '
    '기체의 3D 위치(x, y, z)를 피드백 받아 목표 위치로의 이동 명령을 생성한다.'
)
doc.add_paragraph('■ 수평 위치 제어 (X-Y)')
doc.add_paragraph('  위치 오차 → PID → 목표 자세(Roll, Pitch) + 목표 추진력')
doc.add_paragraph()
doc.add_paragraph('  전진 방향(X) 이동:')
doc.add_paragraph('    X 오차 → PID → 꼬리 진동 주파수/진폭 (추진력 조절)')
doc.add_paragraph()
doc.add_paragraph('  횡방향(Y) 이동:')
doc.add_paragraph('    Y 오차 → PID → 목표 Yaw 변경 → 꼬리 좌우 편향')
doc.add_paragraph()
doc.add_paragraph(
    '  헬륨 드론은 비홀로노믹(Non-holonomic) 시스템이다. '
    '즉, 돌고래처럼 전진 방향으로만 추진력이 발생하며, 측면 이동은 '
    '먼저 Yaw 회전 후 전진하는 방식으로 이루어진다. '
    '이는 멀티로터 드론(홀로노믹 시스템)과의 근본적인 차이점이다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 위치 제어 전략: Pure Pursuit 또는 Carrot Following')
doc.add_paragraph(
    '  경로 위의 전방 목표점(Look-ahead point)을 추종하는 방식으로, '
    '비홀로노믹 시스템에 적합한 경로 추종 알고리즘이다.'
)
doc.add_paragraph()
doc.add_paragraph('  1. 현재 위치에서 경로 상의 최근접점(Closest point) 탐색')
doc.add_paragraph('  2. 최근접점에서 전방 L_d(Look-ahead distance)만큼 떨어진 목표점 설정')
doc.add_paragraph('  3. 목표점을 향한 목표 Yaw 각도 계산')
doc.add_paragraph('  4. Yaw 제어 + 추진력 제어로 경로 추종')
add_ref('[참고] Coulter, R. C. (1992). "Implementation of the pure pursuit path tracking algorithm." '
        'Carnegie Mellon University, CMU-RI-TR-92-01.')

doc.add_heading('6.4 고도 제어 (Altitude Control)', level=2)
doc.add_paragraph(
    '헬륨 드론의 고도 제어는 멀티로터와 근본적으로 다르다. '
    '중립 부력 상태에서는 추력을 가하지 않아도 고도가 유지되며, '
    '고도 변화는 가슴지느러미의 상하 틸트로 미세한 수직 분력을 만들어 제어한다.'
)
doc.add_paragraph()
doc.add_paragraph('  고도 제어 입력:')
doc.add_paragraph('    Z 오차 → PID → 가슴지느러미 공통 틸트 각도')
doc.add_paragraph()
doc.add_paragraph('    상승: 가슴지느러미 하향 틸트 → 하향 양력 반작용으로 상승')
doc.add_paragraph('    하강: 가슴지느러미 상향 틸트 → 상향 양력 반작용으로 하강')
doc.add_paragraph('    유지: 지느러미 수평 (부력 = 중력 상태)')
doc.add_paragraph()
doc.add_paragraph(
    '중립 부력을 정확히 설정하면, 고도 제어에 필요한 에너지가 극히 적다. '
    '이는 끊임없이 모터를 회전시켜야 하는 멀티로터 대비 큰 에너지 이점이다. '
    '다만, 실내 온도 변화(공조)에 따른 헬륨 밀도 변화로 부력이 미세하게 변동할 수 있으며, '
    '이를 보상하기 위한 적분(I) 제어기가 필수적이다.'
)

doc.add_heading('6.5 PID 제어기 설계', level=2)
doc.add_paragraph(
    '각 축별 PID 제어기의 초기 파라미터와 튜닝 전략을 제시한다.'
)
doc.add_paragraph('■ PID 제어 법칙:')
doc.add_paragraph('  u(t) = Kp·e(t) + Ki·∫e(τ)dτ + Kd·(de/dt)')
doc.add_paragraph()
add_table(
    ['제어 축', 'Kp (초기값)', 'Ki (초기값)', 'Kd (초기값)', '출력', '비고'],
    [
        ['Roll', '2.0', '0.1', '0.5', '지느러미 차동 각도', '내부 루프'],
        ['Pitch', '2.0', '0.1', '0.5', '지느러미 공통 각도', '내부 루프'],
        ['Yaw', '1.5', '0.05', '0.3', '꼬리 좌우 편향', '내부 루프'],
        ['X (전진)', '1.0', '0.05', '0.8', '꼬리 주파수/진폭', '외부 루프'],
        ['Y (횡)', '1.0', '0.05', '0.8', '목표 Yaw 변경', '외부 루프'],
        ['Z (고도)', '1.5', '0.2', '0.5', '지느러미 틸트', '적분 중요'],
    ]
)
doc.add_paragraph(
    '이 초기값은 시뮬레이션(Gazebo)에서 1차 튜닝 후, 실기체에서 미세 조정한다. '
    '헬륨 드론은 관성이 크고 응답이 느리므로, D(미분) 게인을 적절히 설정하여 '
    '오버슈트를 방지하는 것이 중요하다.'
)
doc.add_paragraph()
doc.add_paragraph('■ PID 튜닝 순서:')
doc.add_paragraph('  1. 내부 루프(자세)부터 튜닝 — Roll → Pitch → Yaw 순서')
doc.add_paragraph('  2. 자세 안정화 확인 후 외부 루프(위치) 튜닝')
doc.add_paragraph('  3. 고도(Z) 제어는 부력 평형 확인 후 별도 튜닝')
doc.add_paragraph('  4. Ziegler-Nichols 방법 또는 시행착오법 적용')
add_ref('[참고] Åström, K. J., & Murray, R. M. (2021). "Feedback Systems: An Introduction for Scientists '
        'and Engineers." Princeton University Press, 2nd Edition.')

doc.add_heading('6.6 꼬리 진동 추진력 모델링', level=2)
doc.add_paragraph(
    '꼬리지느러미의 진동 운동을 추진력으로 변환하는 수학적 모델을 수립한다.'
)
doc.add_paragraph('■ 꼬리 운동 모델 (단순화)')
doc.add_paragraph('  θ(t) = A · sin(2πft + φ)')
doc.add_paragraph('  여기서:')
doc.add_paragraph('    A = 진폭 (rad), 제어 가능 범위: 0.35~0.70 rad (20°~40°)')
doc.add_paragraph('    f = 주파수 (Hz), 제어 가능 범위: 1~3 Hz')
doc.add_paragraph('    φ = 위상')
doc.add_paragraph()
doc.add_paragraph('■ 순간 추진력 모델')
doc.add_paragraph('  F_thrust(t) = 0.5 × ρ_air × S_fin × [A·2πf·cos(2πft)]² × C_T(St)')
doc.add_paragraph()
doc.add_paragraph('■ 평균 추진력')
doc.add_paragraph('  F̄_thrust = 0.5 × ρ_air × S_fin × (A·2πf)² × 0.5 × C_T(St)')
doc.add_paragraph('           = ρ_air × π² × S_fin × A² × f² × C_T(St)')
doc.add_paragraph()
doc.add_paragraph(
    '여기서 C_T(St)는 Strouhal Number에 의존하는 추력계수이며, '
    'St = f·2A·sin(A_max)/U 이다. St ≈ 0.2~0.4에서 최대 효율을 보인다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 수치 예시 (2m급 돌고래 드론)')
doc.add_paragraph('  S_fin = 0.02 m² (꼬리 면적)')
doc.add_paragraph('  A = 0.52 rad (30°)')
doc.add_paragraph('  f = 2 Hz')
doc.add_paragraph('  ρ_air = 1.225 kg/m³')
doc.add_paragraph('  C_T ≈ 0.3')
doc.add_paragraph('  F̄_thrust ≈ 1.225 × 9.87 × 0.02 × 0.27 × 4 × 0.3 ≈ 0.078 N ≈ 8 gf')
doc.add_paragraph()
doc.add_paragraph(
    '약 8gf의 추진력은 작아 보이지만, 중립 부력 상태의 헬륨 드론(~500g 질량)에게는 '
    '0.16 m/s²의 가속도를 제공하며, 실내 환경에서 약 0.3~0.5 m/s의 순항 속도를 '
    '달성하기에 충분하다.'
)
add_ref('[참고] Lighthill, M. J. (1970). "Aquatic animal propulsion of high hydromechanical efficiency." '
        'Journal of Fluid Mechanics, 44(2), 265-301.')

doc.add_heading('6.7 제어 입력 매핑 (Control Allocation)', level=2)
doc.add_paragraph(
    'PID 제어기의 출력을 실제 서보 명령으로 변환하는 Control Allocation 행렬을 정의한다. '
    '이는 멀티로터의 모터 믹서(Mixer)에 해당한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 제어 변수 → 서보 명령 매핑')
doc.add_paragraph()
doc.add_paragraph('  [δ_tail_pitch]     [1   0   0   0] [u_thrust  ]')
doc.add_paragraph('  [δ_tail_yaw  ]  =  [0   0   0   1] [u_roll    ]')
doc.add_paragraph('  [δ_pec_left  ]     [0   1   1   0] [u_pitch   ]')
doc.add_paragraph('  [δ_pec_right ]     [0   1  -1   0] [u_yaw     ]')
doc.add_paragraph()
doc.add_paragraph('  여기서:')
doc.add_paragraph('    u_thrust = 꼬리 진동 파라미터 (주파수 × 진폭) → 꼬리 상하 서보')
doc.add_paragraph('    u_pitch  = 피치 제어 출력 → 가슴지느러미 공통 각도')
doc.add_paragraph('    u_roll   = 롤 제어 출력 → 가슴지느러미 차동 각도')
doc.add_paragraph('    u_yaw    = 요 제어 출력 → 꼬리 좌우 편향')
doc.add_paragraph()
doc.add_paragraph('  ※ 추진-조향 커플링 보상:')
doc.add_paragraph('    꼬리 상하(추진)와 꼬리 좌우(Yaw)가 물리적으로 결합되어 있으므로,')
doc.add_paragraph('    Yaw 제어 시 추진력이 변동할 수 있다. 이를 보상하기 위해:')
doc.add_paragraph('    • u_thrust에 Yaw 편향 각도에 따른 보상항 추가')
doc.add_paragraph('    • Simulink에서 커플링 전달함수를 모델링하여 피드포워드 보상 적용')
doc.add_paragraph()
doc.add_paragraph(
    '각 서보의 출력은 물리적 한계(±45°)로 클리핑(Clipping)하며, '
    '우선순위는 자세 안정화 > 고도 > 수평 이동 순으로 설정한다. '
    '제어 입력이 서보 범위를 초과할 경우, 우선순위가 낮은 축의 제어량을 줄인다.'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 7장. 군집 비행 제어
# ═══════════════════════════════════════════
doc.add_heading('7장. 군집 비행 제어', level=1)

doc.add_heading('7.1 군집 제어 아키텍처', level=2)
doc.add_paragraph(
    '5~10대의 돌고래 헬륨 드론을 동시에 제어하기 위한 군집 아키텍처를 설계한다.'
)
doc.add_paragraph('■ 중앙 집중형 (Centralized) 아키텍처 채택')
doc.add_paragraph()
doc.add_paragraph(
    '아트 퍼포먼스의 특성상 사전 계획된 안무(Choreography)를 정확히 재현해야 하므로, '
    '중앙 집중형 제어가 적합하다. 분산형(Decentralized) 군집 제어는 '
    '창발적(Emergent) 행동에 유리하지만, 정확한 포메이션 재현에는 불리하다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 소프트웨어 스택')
doc.add_paragraph('  Ground Station (PC):')
doc.add_paragraph('    • 안무 재생 엔진 (Python/C++)')
doc.add_paragraph('    • 모션캡처 인터페이스 (NatNet SDK 또는 OpenCV)')
doc.add_paragraph('    • 다중 드론 위치 제어기 (각 드론별 독립 PID)')
doc.add_paragraph('    • 충돌 회피 모듈')
doc.add_paragraph('    • WiFi 통신 매니저 (UDP 브로드캐스트)')
doc.add_paragraph()
doc.add_paragraph('  각 드론 (ESP32):')
doc.add_paragraph('    • WiFi 수신 (목표 자세/추진력)')
doc.add_paragraph('    • IMU 기반 자세 안정화')
doc.add_paragraph('    • 서보 PWM 출력')
doc.add_paragraph('    • LED 색상/밝기 제어')
doc.add_paragraph('    • 상태 보고 (배터리, IMU 데이터)')
add_ref('[참고] Preiss, J. A., et al. (2017). "Crazyswarm: A large nano-quadcopter swarm." '
        'IEEE ICRA 2017. — 중앙 집중형 군집 제어의 대표적 구현 사례')

doc.add_heading('7.2 경로 계획 및 안무(Choreography) 설계', level=2)
doc.add_paragraph(
    '군집 아트 퍼포먼스의 핵심은 안무 데이터의 설계와 재생이다.'
)
doc.add_paragraph('■ 안무 데이터 형식')
doc.add_paragraph('  각 드론의 시간별 목표 위치를 정의하는 Keyframe 기반 시스템:')
doc.add_paragraph()
doc.add_paragraph('  {')
doc.add_paragraph('    "drone_id": 1,')
doc.add_paragraph('    "keyframes": [')
doc.add_paragraph('      {"t": 0.0, "x": 0.0, "y": 0.0, "z": 1.5, "led": [0,0,255]},')
doc.add_paragraph('      {"t": 5.0, "x": 2.0, "y": 1.0, "z": 2.0, "led": [0,255,0]},')
doc.add_paragraph('      {"t": 10.0, "x": -1.0, "y": 2.0, "z": 1.0, "led": [255,0,0]}')
doc.add_paragraph('    ]')
doc.add_paragraph('  }')
doc.add_paragraph()
doc.add_paragraph('■ 궤적 생성: Minimum Snap Trajectory')
doc.add_paragraph(
    '  Keyframe 사이의 궤적은 Minimum Snap(최소 스냅) 다항식으로 보간한다. '
    '이 방법은 Mellinger & Kumar(2011)가 제안한 것으로, '
    '부드럽고 에너지 효율적인 궤적을 생성한다.'
)
doc.add_paragraph()
doc.add_paragraph('  궤적 다항식: p(t) = c₀ + c₁t + c₂t² + c₃t³ + c₄t⁴ + c₅t⁵ + c₆t⁶ + c₇t⁷')
doc.add_paragraph('  최적화 목표: minimize ∫(d⁴p/dt⁴)² dt (Snap의 제곱 적분 최소화)')
doc.add_paragraph()
doc.add_paragraph('■ 안무 설계 도구')
doc.add_paragraph('  • 3D GUI 기반 안무 편집기 (향후 개발)')
doc.add_paragraph('  • Blender 애니메이션 → 궤적 데이터 변환 (기존 도구 활용 가능)')
doc.add_paragraph('  • CSV/JSON 기반 수동 편집')
doc.add_paragraph('  • 음악 싱크: 비트 감지 → 자동 키프레임 생성')
add_ref('[참고] Mellinger, D., & Kumar, V. (2011). "Minimum snap trajectory generation and control '
        'for quadrotors." IEEE ICRA 2011.')

doc.add_heading('7.3 충돌 회피 알고리즘', level=2)
doc.add_paragraph(
    '다수의 드론이 동시에 비행할 때, 충돌 회피는 안전성의 핵심이다.'
)
doc.add_paragraph('■ 2단계 충돌 회피 전략')
doc.add_paragraph()
doc.add_paragraph('  1단계: 사전 충돌 검사 (Offline)')
doc.add_paragraph('    안무 데이터 생성 시점에 전체 궤적을 분석하여 충돌 가능 구간을 사전 검출.')
doc.add_paragraph('    최소 안전 거리: D_safe = 0.8m (2m급 기체의 최대 폭 0.5m + 여유 0.3m)')
doc.add_paragraph()
doc.add_paragraph('  2단계: 실시간 충돌 회피 (Online)')
doc.add_paragraph('    비행 중 예기치 않은 근접 상황 발생 시 실시간 회피 기동 실행.')
doc.add_paragraph('    알고리즘: ORCA (Optimal Reciprocal Collision Avoidance)')
doc.add_paragraph()
doc.add_paragraph(
    'ORCA는 Van den Berg et al.(2011)이 제안한 속도 장애물(Velocity Obstacle) 기반 '
    '알고리즘으로, 각 드론이 독립적으로 회피 속도를 계산하되 '
    '상호 대칭적인 회피 행동을 보장한다. 10대 규모에서도 실시간 연산이 가능하다.'
)
add_ref('[참고] Van den Berg, J., et al. (2011). "Reciprocal n-body collision avoidance." '
        'Robotics Research, Springer, 3-19.')

doc.add_heading('7.4 동기화 및 포메이션 제어', level=2)
doc.add_paragraph(
    '아트 퍼포먼스에서 시각적 효과를 극대화하려면 드론 간의 시간 동기화와 '
    '포메이션(대형) 정확도가 중요하다.'
)
doc.add_paragraph('■ 시간 동기화')
doc.add_paragraph('  • NTP(Network Time Protocol) 기반 시간 동기화: ±1ms 정밀도')
doc.add_paragraph('  • Ground Station의 마스터 클럭에 전 드론이 동기화')
doc.add_paragraph('  • 안무 재생의 시간 기준점(Epoch)을 공유')
doc.add_paragraph()
doc.add_paragraph('■ 포메이션 패턴 예시')
doc.add_paragraph('  • V자 편대 (돌고래 무리 유영 모방)')
doc.add_paragraph('  • 원형 선회 (나선형 상승/하강)')
doc.add_paragraph('  • 분산 → 수렴 (폭발/집합 효과)')
doc.add_paragraph('  • 파동(Wave) 패턴 (연속적 높이 변화)')
doc.add_paragraph('  • 하트/별 등 2D 형상 (관객 방향에서 관찰)')
doc.add_paragraph()
doc.add_paragraph(
    '각 포메이션은 가상 구조체(Virtual Structure) 방식으로 구현한다. '
    '가상 강체(Virtual Rigid Body)의 이동/회전에 따라 각 드론의 목표 위치가 결정되며, '
    '이를 통해 전체 포메이션의 이동, 회전, 크기 변환을 간단하게 표현할 수 있다.'
)
add_ref('[참고] Ren, W., & Beard, R. W. (2008). "Distributed Consensus in Multi-vehicle Cooperative Control." '
        'Springer London.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 8장. Simulink 모델링 및 제어기 설계
# ═══════════════════════════════════════════
doc.add_heading('8장. Simulink 모델링 및 제어기 설계', level=1)

doc.add_paragraph(
    '본 장에서는 MATLAB/Simulink를 이용한 돌고래 헬륨 드론의 플랜트 모델링과 '
    '비행 제어기 설계 과정을 상세히 기술한다. Simulink에서 제어기를 설계·검증한 후, '
    'Embedded Coder를 통해 ESP32용 C 코드를 생성하거나, 확정된 파라미터를 수동으로 '
    'C++에 포팅하는 워크플로우를 따른다.'
)

doc.add_heading('8.1 Simulink 전체 모델 구조', level=2)
doc.add_paragraph(
    'Simulink 모델은 크게 4개의 서브시스템으로 구성된다.'
)
doc.add_paragraph()
doc.add_paragraph('■ Top-Level 모델 구조')
doc.add_paragraph()
doc.add_paragraph('  ┌─────────────────────────────────────────────────────────┐')
doc.add_paragraph('  │                    Simulink Top Model                    │')
doc.add_paragraph('  │                                                         │')
doc.add_paragraph('  │  [Trajectory Gen] → [Position Ctrl] → [Attitude Ctrl]  │')
doc.add_paragraph('  │        │                  │                  │          │')
doc.add_paragraph('  │        │                  │          [Control Alloc]    │')
doc.add_paragraph('  │        │                  │                  │          │')
doc.add_paragraph('  │        │                  │          [Plant Model]      │')
doc.add_paragraph('  │        │                  │            ↓     ↑         │')
doc.add_paragraph('  │        │                  └──── Position ────┘         │')
doc.add_paragraph('  │        │                        Attitude ──→ Att Ctrl  │')
doc.add_paragraph('  └─────────────────────────────────────────────────────────┘')
doc.add_paragraph()
doc.add_paragraph('■ 서브시스템 목록')
doc.add_paragraph()
doc.add_paragraph('  1. Trajectory Generator (궤적 생성기)')
doc.add_paragraph('     - 입력: Waypoint 배열 (x, y, z, t)')
doc.add_paragraph('     - 출력: 시간별 목표 위치/속도')
doc.add_paragraph('     - Minimum Snap 다항식 보간')
doc.add_paragraph()
doc.add_paragraph('  2. Position Controller (위치 제어기) — 외부 루프')
doc.add_paragraph('     - 입력: 목표 위치, 현재 위치 (Plant 피드백)')
doc.add_paragraph('     - 출력: 목표 자세 (Roll, Pitch, Yaw) + 목표 추진력')
doc.add_paragraph('     - PID 또는 MPC')
doc.add_paragraph('     - 실행 주기: 10~20 Hz (Rate Transition 블록)')
doc.add_paragraph()
doc.add_paragraph('  3. Attitude Controller (자세 제어기) — 내부 루프')
doc.add_paragraph('     - 입력: 목표 자세, 현재 자세 (IMU 피드백)')
doc.add_paragraph('     - 출력: 제어 명령 (u_thrust, u_roll, u_pitch, u_yaw)')
doc.add_paragraph('     - PID')
doc.add_paragraph('     - 실행 주기: 50~100 Hz')
doc.add_paragraph()
doc.add_paragraph('  4. Plant Model (플랜트 모델) — 돌고래 드론 물리')
doc.add_paragraph('     - 입력: 서보 명령 4채널')
doc.add_paragraph('     - 출력: 위치 (x,y,z), 자세 (Roll,Pitch,Yaw), 속도')
doc.add_paragraph('     - 6-DOF 운동방정식 + 부력 + 추진 + 항력')
doc.add_paragraph()
add_ref('[참고] MathWorks. "Aerospace Blockset — 6DOF (Euler Angles)." Documentation.')

doc.add_heading('8.2 플랜트 모델 (Dolphin Drone Plant)', level=2)
doc.add_paragraph(
    '플랜트 모델은 돌고래 헬륨 드론의 물리적 동역학을 시뮬레이션한다. '
    '핵심은 부력, 핀 추진력, 공기 항력, 중력의 합력에 의한 6-DOF 운동이다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 6-DOF 운동 방정식')
doc.add_paragraph('  병진 운동: m × a = F_buoyancy + F_gravity + F_thrust + F_drag')
doc.add_paragraph('  회전 운동: J × α = τ_fins + τ_tail + τ_dorsal_passive + τ_damping')
doc.add_paragraph()
doc.add_paragraph('■ Simulink 구현: 서브시스템 내부 블록')
doc.add_paragraph()
doc.add_paragraph('  (a) 부력 블록 (Buoyancy)')
doc.add_paragraph('      F_b = [0; 0; (ρ_air - ρ_He) × V × g]  (상향)')
doc.add_paragraph('      τ_b = r_CB × F_b  (CB-CG 오프셋에 의한 복원 토크)')
doc.add_paragraph()
doc.add_paragraph('  (b) 중력 블록 (Gravity)')
doc.add_paragraph('      F_g = [0; 0; -m × g]  (하향)')
doc.add_paragraph()
doc.add_paragraph('  (c) 꼬리 추진력 블록 (Tail Thrust)')
doc.add_paragraph('      입력: 진동 주파수(f), 진폭(A), 좌우 편향(δ_yaw)')
doc.add_paragraph('      θ(t) = A × sin(2πft)')
doc.add_paragraph('      F_thrust_mag = ρ_air × π² × S_fin × A² × f² × C_T')
doc.add_paragraph('      F_thrust = F_thrust_mag × [cos(δ_yaw); sin(δ_yaw); 0]  (기체 좌표계)')
doc.add_paragraph('      τ_yaw = r_tail × F_thrust × sin(δ_yaw)  (Yaw 모멘트)')
doc.add_paragraph('      ※ 커플링 모델: Yaw 편향 시 전진 추력 감소 = F × cos(δ_yaw)')
doc.add_paragraph()
doc.add_paragraph('  (d) 가슴지느러미 블록 (Pectoral Fins)')
doc.add_paragraph('      입력: δ_left, δ_right (각 지느러미 각도)')
doc.add_paragraph('      F_pec = 0.5 × ρ_air × V² × S_pec × C_L(δ) (속도 의존)')
doc.add_paragraph('      τ_roll = (F_pec_L - F_pec_R) × r_pec  (Roll 모멘트)')
doc.add_paragraph('      τ_pitch = (F_pec_L + F_pec_R) × r_pec_x  (Pitch 모멘트)')
doc.add_paragraph('      F_z = (F_pec_L + F_pec_R) × sin(δ_common)  (수직 분력 → 고도)')
doc.add_paragraph()
doc.add_paragraph('  (e) 등지느러미 수동 감쇠 블록 (Dorsal Fin Passive)')
doc.add_paragraph('      τ_dorsal_yaw = -C_dorsal × ψ̇  (Yaw rate 감쇠)')
doc.add_paragraph('      τ_dorsal_roll = -C_dorsal_r × φ̇  (Roll rate 감쇠)')
doc.add_paragraph('      ※ 속도가 높을수록 감쇠 효과 증가')
doc.add_paragraph()
doc.add_paragraph('  (f) 공기 항력 블록 (Air Drag)')
doc.add_paragraph('      F_drag = -0.5 × ρ_air × Cd × S_ref × |V| × V')
doc.add_paragraph('      Cd ≈ 0.30 (엔벨로프+지느러미 등가 항력계수)')
doc.add_paragraph()
doc.add_paragraph('  (g) 6-DOF 적분 블록')
doc.add_paragraph('      Aerospace Blockset의 "6DOF (Euler Angles)" 블록 사용')
doc.add_paragraph('      또는 Custom MATLAB Function으로 RK4 적분 구현')
doc.add_paragraph()
doc.add_paragraph('■ MATLAB 파라미터 스크립트 (init_params.m)')
doc.add_paragraph()
doc.add_paragraph('  % 돌고래 드론 물리 파라미터')
doc.add_paragraph('  params.m = 0.279;          % 총 질량 [kg] (음성부력 -1g)')
doc.add_paragraph('  params.V_envelope = 0.262;  % 엔벨로프 체적 [m³]')
doc.add_paragraph('  params.rho_air = 1.225;     % 공기 밀도 [kg/m³]')
doc.add_paragraph('  params.rho_He = 0.164;      % 헬륨 밀도 [kg/m³]')
doc.add_paragraph('  params.g = 9.81;            % 중력 가속도 [m/s²]')
doc.add_paragraph('  params.S_tail = 0.02;       % 꼬리 면적 [m²]')
doc.add_paragraph('  params.S_pec = 0.01;        % 가슴지느러미 면적 [m²]')
doc.add_paragraph('  params.S_dorsal = 0.015;    % 등지느러미 면적 [m²]')
doc.add_paragraph('  params.Cd = 0.04;           % 형상 항력계수')
doc.add_paragraph('  params.S_ref = 0.15;        % 기준 면적 [m²]')
doc.add_paragraph('  params.C_T = 0.3;           % 꼬리 추력계수')
doc.add_paragraph('  params.r_tail = [-0.8; 0; 0]; % 꼬리 위치 [m]')
doc.add_paragraph('  params.r_pec = [0.1; 0.25; 0]; % 가슴지느러미 위치 [m]')
doc.add_paragraph('  params.r_CB = [0; 0; 0.05];    % 부력중심 오프셋 [m]')
doc.add_paragraph('  params.J = diag([0.01, 0.08, 0.08]); % 관성 모멘트 [kg·m²]')
doc.add_paragraph('  params.C_dorsal_yaw = 0.005; % 등지느러미 Yaw 감쇠계수')
doc.add_paragraph('  params.C_dorsal_roll = 0.003; % 등지느러미 Roll 감쇠계수')
doc.add_paragraph()
add_ref('[참고] Li, Y., et al. (2011). "Dynamic modeling and analysis of a tri-lobed airship." '
        'Aerospace Science and Technology, 15(6), 481-488.')

doc.add_heading('8.3 자세 제어기 서브시스템', level=2)
doc.add_paragraph(
    '자세 제어기는 ESP32에 탑재될 내부 루프로, Simulink에서 설계 후 코드 생성한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ Simulink 블록 구성')
doc.add_paragraph()
doc.add_paragraph('  [목표 자세] ──→ ⊕ ──→ [PID Roll]  ──→ u_roll')
doc.add_paragraph('                 ↑-       ')
doc.add_paragraph('  [IMU 피드백] ──┘')
doc.add_paragraph()
doc.add_paragraph('  동일 구조로 Pitch, Yaw 각각 독립 PID')
doc.add_paragraph()
doc.add_paragraph('■ 각 축 PID 블록 구성')
doc.add_paragraph('  • Simulink PID Controller 블록 사용 (Continuous 또는 Discrete)')
doc.add_paragraph('  • Anti-Windup: Clamping 방식 (적분 포화 방지)')
doc.add_paragraph('  • Derivative Filter: N = 100 (미분항 노이즈 필터)')
doc.add_paragraph('  • Output Saturation: ±1.0 (정규화 출력)')
doc.add_paragraph()
doc.add_paragraph('■ Madgwick 필터 블록')
doc.add_paragraph('  • MATLAB Function 블록으로 구현')
doc.add_paragraph('  • 입력: 가속도(ax,ay,az), 자이로(gx,gy,gz)')
doc.add_paragraph('  • 출력: 쿼터니언 → 오일러각 변환')
doc.add_paragraph('  • 파라미터: β = 0.1 (초기값)')
doc.add_paragraph()
doc.add_paragraph('■ PID 자동 튜닝')
doc.add_paragraph('  Simulink의 PID Tuner 앱을 활용하여 각 축을 개별 튜닝:')
doc.add_paragraph('  1. Plant 모델에서 해당 축의 Open-Loop 전달함수 추출')
doc.add_paragraph('  2. PID Tuner에서 응답 시간(Response Time)과 과도 응답(Transient) 조절')
doc.add_paragraph('  3. Bandwidth: 1~5 Hz (헬륨 드론의 느린 동역학에 맞춤)')
doc.add_paragraph('  4. Phase Margin: 60° 이상 확보')
doc.add_paragraph()
doc.add_paragraph('  % MATLAB 코드 예시: Roll축 PID 자동 튜닝')
doc.add_paragraph('  G_roll = linearize(mdl, "Attitude_Controller/Roll");')
doc.add_paragraph('  pidTuner(G_roll, "PID");  % GUI 기반 튜닝')
doc.add_paragraph()
add_ref('[참고] MathWorks. "PID Tuner — Tune PID Controller." Documentation.')

doc.add_heading('8.4 위치 제어기 서브시스템', level=2)
doc.add_paragraph(
    '위치 제어기는 Ground Station PC에서 실행되는 외부 루프이다. '
    'Simulink에서 설계하고, 확정된 파라미터를 Python/C++ Ground Station 코드에 포팅한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 비홀로노믹 시스템 고려')
doc.add_paragraph('  돌고래 드론은 전진 방향으로만 추진력이 발생하므로,')
doc.add_paragraph('  일반적인 XYZ 독립 PID가 아닌 극좌표 기반 제어를 사용한다:')
doc.add_paragraph()
doc.add_paragraph('  1. 목표점까지의 거리(d)와 방위각(ψ_target) 계산')
doc.add_paragraph('  2. Yaw 제어: 현재 ψ → 목표 ψ_target으로 회전')
doc.add_paragraph('  3. 전진 제어: d가 줄어들도록 꼬리 추진력 조절')
doc.add_paragraph('  4. 고도 제어: Z 오차 → 가슴지느러미 틸트')
doc.add_paragraph()
doc.add_paragraph('■ Simulink 블록 구성')
doc.add_paragraph()
doc.add_paragraph('  [목표 위치] ──→ [Polar Error] ──→ [Yaw PID] ──→ 목표 Yaw')
doc.add_paragraph('  [현재 위치] ──→ [Calculator ] ──→ [Dist PID] ──→ 목표 추진력')
doc.add_paragraph('                                ──→ [Alt PID]  ──→ 목표 Pitch')
doc.add_paragraph()
doc.add_paragraph('■ Rate Transition')
doc.add_paragraph('  위치 제어기(10Hz)와 자세 제어기(100Hz)의 샘플링 레이트가 다르므로,')
doc.add_paragraph('  Simulink Rate Transition 블록으로 멀티레이트 시스템을 구현한다.')
doc.add_paragraph()
add_ref('[참고] Mellinger, D., & Kumar, V. (2011). "Minimum snap trajectory generation and control '
        'for quadrotors." IEEE ICRA 2011.')

doc.add_heading('8.5 커플링 보상 및 Control Allocation', level=2)
doc.add_paragraph(
    '꼬리가 추진(상하)과 Yaw(좌우)를 동시에 수행하므로 발생하는 '
    '추진-조향 커플링을 Simulink에서 모델링하고 보상한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 커플링 모델')
doc.add_paragraph('  꼬리 좌우 편향(δ_yaw) 시:')
doc.add_paragraph('    실효 전진 추력: F_forward = F_thrust × cos(δ_yaw)')
doc.add_paragraph('    Yaw 추력:      F_yaw = F_thrust × sin(δ_yaw)')
doc.add_paragraph('    → δ_yaw = 30° 일 때 전진 추력 ~13% 감소')
doc.add_paragraph()
doc.add_paragraph('■ 피드포워드 보상 블록')
doc.add_paragraph('  Simulink MATLAB Function 블록으로 구현:')
doc.add_paragraph()
doc.add_paragraph('  function [f_comp, A_comp] = coupling_compensator(f, A, delta_yaw)')
doc.add_paragraph('      % 추진력 보상: Yaw 편향에 따른 전진력 감소를 주파수/진폭 증가로 보상')
doc.add_paragraph('      thrust_loss = 1 / cos(delta_yaw);')
doc.add_paragraph('      f_comp = f * sqrt(thrust_loss);    % 주파수 보상')
doc.add_paragraph('      A_comp = A * sqrt(thrust_loss);    % 진폭 보상')
doc.add_paragraph('      % Saturation')
doc.add_paragraph('      f_comp = min(f_comp, 3.0);  % 최대 3Hz')
doc.add_paragraph('      A_comp = min(A_comp, 0.70); % 최대 40°')
doc.add_paragraph('  end')
doc.add_paragraph()
doc.add_paragraph('■ Control Allocation 블록')
doc.add_paragraph('  제어 명령 → 서보 명령 변환을 Simulink Subsystem으로 구현:')
doc.add_paragraph()
doc.add_paragraph('  입력: u_thrust, u_roll, u_pitch, u_yaw (정규화 ±1)')
doc.add_paragraph('  출력: δ_tail_pitch(주파수/진폭), δ_tail_yaw, δ_pec_L, δ_pec_R')
doc.add_paragraph()
doc.add_paragraph('  δ_tail_freq  = u_thrust × (f_max - f_min) + f_min')
doc.add_paragraph('  δ_tail_amp   = |u_thrust| × (A_max - A_min) + A_min')
doc.add_paragraph('  δ_tail_yaw   = u_yaw × δ_yaw_max')
doc.add_paragraph('  δ_pec_L      = u_pitch × δ_pitch_max + u_roll × δ_roll_max')
doc.add_paragraph('  δ_pec_R      = u_pitch × δ_pitch_max - u_roll × δ_roll_max')
doc.add_paragraph()
doc.add_paragraph('  각 출력은 Saturation 블록으로 물리적 한계 내로 클리핑')
doc.add_paragraph('  우선순위: 자세 안정화 > 고도 > 수평 이동 (Priority Allocator)')

doc.add_heading('8.6 시뮬레이션 시나리오 및 튜닝 절차', level=2)
doc.add_paragraph(
    'Simulink 모델 완성 후 다음 순서로 시뮬레이션과 튜닝을 진행한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ Step 1: 오픈루프 테스트 (Plant 검증)')
doc.add_paragraph('  • 서보 입력을 수동으로 인가하여 Plant 응답 확인')
doc.add_paragraph('  • 부력 평형 확인 (정지 상태에서 고도 유지)')
doc.add_paragraph('  • 꼬리 진동 시 전진 속도 프로파일 확인')
doc.add_paragraph('  • 등지느러미 수동 감쇠 효과 확인')
doc.add_paragraph()
doc.add_paragraph('■ Step 2: 내부 루프 튜닝 (자세 제어)')
doc.add_paragraph('  • Roll → Pitch → Yaw 순서로 개별 튜닝')
doc.add_paragraph('  • PID Tuner 앱으로 초기 게인 산출')
doc.add_paragraph('  • Step 응답 확인: 오버슈트 < 10%, 정정시간 < 3초')
doc.add_paragraph('  • Bode Plot으로 안정성 여유 확인 (PM > 60°, GM > 6dB)')
doc.add_paragraph()
doc.add_paragraph('■ Step 3: 외부 루프 튜닝 (위치 제어)')
doc.add_paragraph('  • 내부 루프 폐루프 안정 확인 후 진행')
doc.add_paragraph('  • Waypoint 추종 시뮬레이션')
doc.add_paragraph('  • 위치 오차 < 5cm (정상 상태)')
doc.add_paragraph('  • 외란(0.2 m/s 기류) 인가 시 복귀 성능 확인')
doc.add_paragraph()
doc.add_paragraph('■ Step 4: 커플링 보상 검증')
doc.add_paragraph('  • Yaw 30° 회전 시 전진 속도 변동 측정')
doc.add_paragraph('  • 보상 ON/OFF 비교로 효과 확인')
doc.add_paragraph('  • 동시 전진 + Yaw 회전 시나리오')
doc.add_paragraph()
doc.add_paragraph('■ Step 5: 다중 드론 시뮬레이션')
doc.add_paragraph('  • Plant 모델을 N개 복제 (Model Reference 활용)')
doc.add_paragraph('  • 군집 제어 서버 + N개 위치/자세 제어기')
doc.add_paragraph('  • 포메이션 비행 및 충돌 회피 검증')
doc.add_paragraph()
doc.add_paragraph('■ Step 6: 코드 생성 및 배포')
doc.add_paragraph('  • Simulink Coder / Embedded Coder로 자세 제어기 C 코드 생성')
doc.add_paragraph('  • ESP32 Arduino 프레임워크에 통합')
doc.add_paragraph('  • 또는 확정된 PID 게인을 C++ 코드에 수동 포팅')
doc.add_paragraph()
doc.add_paragraph('■ 확정된 자세 제어기 게인 (Simulink 검증 완료)')
add_table(
    ['축', 'Outer P', 'Inner P', 'Inner I', 'Inner D', '비고'],
    [
        ['Roll (p)', '1.4', '0.30', '0.01', '0.10', 'Roll 1~3deg가 물리적 한계'],
        ['Pitch (q)', '0.25', '0.20', '0.00', '0.02', '꼬리 추진 커플링으로 ±2deg 진동(감쇠)'],
        ['Yaw (r)', '1.0', '0.20', '0.02', '0.03', 'Yaw가 주 조향축'],
    ]
)
doc.add_paragraph('■ RC 모드 검증 결과')
doc.add_paragraph('  • Manual Direct: RC 스틱 → 서보 직접 제어 (PID 우회)')
doc.add_paragraph('  • Stabilized: RC 스틱 → 자세 PID 보조')
doc.add_paragraph('  • Mode Switch: CH5 >= 0.5 → Stabilized, 아니면 Manual')
doc.add_paragraph('  • Deadband ±0.05, RC_Valid failsafe, Servo dynamics 포함')
doc.add_paragraph('  • 부력 설정: 음성부력 -1g (fail-safe 자연 착륙)')
doc.add_paragraph('  • Roll 물리적 한계: 1~3deg (CB-CG 펜듈럼 복원 > 지느러미 토크)')
doc.add_paragraph('  • 조종 방식: Yaw(꼬리)로 방향전환, Roll은 미세 보조')
doc.add_paragraph()
doc.add_paragraph('■ 필요 MATLAB 툴박스')
add_table(
    ['툴박스', '용도', '필수 여부'],
    [
        ['Simulink', '모델링/시뮬레이션', '필수'],
        ['Control System Toolbox', 'PID 튜닝, Bode Plot, 안정성 분석', '필수'],
        ['Aerospace Blockset', '6-DOF 블록, 좌표 변환', '권장'],
        ['Simulink Coder', 'C 코드 자동 생성', '선택 (수동 포팅 가능)'],
        ['Embedded Coder', 'ESP32 타겟 최적화 코드 생성', '선택'],
        ['Robotics System Toolbox', 'ROS2 연동 (Gazebo 연결 시)', '선택'],
        ['Simulink 3D Animation', '3D 비행 시각화', '선택'],
    ]
)
add_ref('[참고] MathWorks. "Simulink Coder — Generate C and C++ Code from Simulink Models." Documentation.')
add_ref('[참고] Åström, K. J., & Murray, R. M. (2021). "Feedback Systems." Princeton University Press.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 9장. Gazebo 연동 시뮬레이션
# ═══════════════════════════════════════════
doc.add_heading('9장. Gazebo 연동 시뮬레이션', level=1)

doc.add_heading('9.1 시뮬레이션 환경 구축', level=2)
doc.add_paragraph(
    '실제 기체 제작 전에 Gazebo 시뮬레이터에서 비행 제어 알고리즘을 검증한다.'
)
doc.add_paragraph('■ 시뮬레이션 환경 구성')
doc.add_paragraph('  • 시뮬레이터: Gazebo Harmonic (ROS2 Humble/Iron 호환)')
doc.add_paragraph('  • 물리 엔진: DART 또는 ODE')
doc.add_paragraph('  • 환경: 실내 공간 모델 (10m × 10m × 5m)')
doc.add_paragraph('  • 드론 모델: 돌고래 형상 URDF/SDF')
doc.add_paragraph()
doc.add_paragraph('■ 필요한 커스텀 플러그인')
doc.add_paragraph('  1. 부력 플러그인 (Buoyancy Plugin): 헬륨 부력 시뮬레이션')
doc.add_paragraph('  2. 핀 추진 플러그인 (Fin Propulsion Plugin): 꼬리/지느러미 추력 모델')
doc.add_paragraph('  3. 공기 저항 플러그인 (Air Drag Plugin): 형상 항력 시뮬레이션')

doc.add_heading('9.2 부력 물리 모델링', level=2)
doc.add_paragraph(
    'Gazebo의 기본 물리 엔진에는 공기 부력 모델이 없으므로, '
    '커스텀 부력 플러그인을 개발해야 한다.'
)
doc.add_paragraph('■ 부력 플러그인 동작 원리')
doc.add_paragraph('  매 시뮬레이션 스텝마다:')
doc.add_paragraph('  1. 기체의 현재 체적(V)과 위치(고도)를 확인')
doc.add_paragraph('  2. 해당 고도에서의 공기 밀도(ρ_air) 계산 (실내이므로 일정)')
doc.add_paragraph('  3. 부력 F = (ρ_air - ρ_helium) × V × g 를 상향으로 적용')
doc.add_paragraph('  4. 부력 적용점(Center of Buoyancy)은 기하학적 중심')
doc.add_paragraph()
doc.add_paragraph(
    '부력 중심(CB)과 무게 중심(CG)의 상대 위치가 정적 안정성을 결정한다. '
    'CB가 CG보다 위에 있으면 정적으로 안정(Self-righting)하며, '
    '이는 종(Pendulum) 효과라 불린다. 돌고래 드론에서는 전자부를 하부에 배치하여 '
    'CB > CG 조건을 확보한다.'
)

doc.add_heading('9.3 Simulink-Gazebo-ROS2 연동', level=2)
doc.add_paragraph(
    'Simulink에서 설계한 제어기를 Gazebo 물리 시뮬레이션과 연동하여 '
    '보다 현실적인 환경에서 검증한다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 연동 구조')
doc.add_paragraph('  Simulink ←→ ROS2 ←→ Gazebo')
doc.add_paragraph()
doc.add_paragraph('  • Simulink: Robotics System Toolbox의 ROS2 블록 사용')
doc.add_paragraph('  • ROS2 토픽: /cmd_servo (Simulink → Gazebo), /imu, /pose (Gazebo → Simulink)')
doc.add_paragraph('  • Gazebo: 커스텀 부력/추진 플러그인이 서보 명령을 받아 물리 시뮬레이션')
doc.add_paragraph()
doc.add_paragraph('  이를 통해 Simulink의 제어기 로직은 그대로 유지하면서,')
doc.add_paragraph('  Plant 모델만 Gazebo의 고정밀 물리 엔진으로 대체할 수 있다.')
doc.add_paragraph()

doc.add_heading('9.4 SITL 기반 비행제어 검증', level=2)
doc.add_paragraph(
    'Software-In-The-Loop(SITL) 방식으로 실제 ESP32 코드를 시뮬레이션 환경에서 검증한다.'
)
doc.add_paragraph('■ SITL 구성')
doc.add_paragraph('  • Gazebo: 물리 시뮬레이션 (부력, 추력, 항력)')
doc.add_paragraph('  • ROS2 노드: 위치 제어기 (PC 역할 대체)')
doc.add_paragraph('  • ESP32 에뮬레이터 또는 PC 포팅: 자세 제어기')
doc.add_paragraph('  • 가상 모션캡처: Gazebo ground truth → 위치 피드백')
doc.add_paragraph()
doc.add_paragraph('■ 검증 시나리오')
doc.add_paragraph('  1. 단일 드론 자세 안정화 테스트')
doc.add_paragraph('  2. 단일 드론 위치 이동 테스트 (Waypoint 추종)')
doc.add_paragraph('  3. 단일 드론 고도 제어 테스트')
doc.add_paragraph('  4. 다중 드론 포메이션 비행 테스트')
doc.add_paragraph('  5. 충돌 회피 시나리오 테스트')
doc.add_paragraph('  6. 전체 안무 시퀀스 리허설')

doc.add_page_break()

# ═══════════════════════════════════════════
# 10장. 설계 사양 및 제작 로드맵
# ═══════════════════════════════════════════
doc.add_heading('10장. 설계 사양 및 제작 로드맵', level=1)

doc.add_heading('10.1 목표 사양', level=2)
add_table(
    ['항목', '사양', '비고'],
    [
        ['전체 길이', '약 2.0m', '돌고래 비율'],
        ['최대 직경', '약 0.50m', 'Fineness Ratio 4:1'],
        ['엔벨로프 체적', '약 262L (0.262m³)', 'Prolate spheroid'],
        ['총 부력', '약 278g', ''],
        ['총 중량 (기체)', '약 279g', '부력+1g 음성부력'],
        ['부력 설정', '음성부력 -1g', 'fail-safe 자연 착륙'],
        ['최대 속도', '약 0.5 m/s', '실내 안전 속도'],
        ['비행 시간', '약 3~6시간', '구성에 따라'],
        ['위치 정밀도', '±5cm 이내', '외부 트래킹 기준'],
        ['군집 규모', '5~10대', ''],
        ['통신', 'WiFi (ESP32)', 'UDP, 10~20Hz'],
        ['제어 방식', '2단 계층 PID', '커스텀 FC'],
    ]
)

doc.add_heading('10.2 BOM 초안 (1대 기준)', level=2)
add_table(
    ['분류', '부품', '수량', '단가(예상)', '소계', '무게'],
    [
        ['외피', 'Mylar 필름 20μm', '1세트', '₩15,000', '₩15,000', '~100g'],
        ['구동', '서보 SG51R (3.5g)', '4', '₩5,000', '₩20,000', '~14g'],
        ['전자', 'ESP32-C3 Mini', '1', '₩8,000', '₩8,000', '~2g'],
        ['전자', 'MPU6050 모듈', '1', '₩3,000', '₩3,000', '~1g'],
        ['전원', '1S LiPo 300mAh', '1', '₩8,000', '₩8,000', '~8g'],
        ['조명', 'WS2812B LED 5개', '1세트', '₩3,000', '₩3,000', '~3g'],
        ['트래킹', 'IR 반사 마커', '4', '₩500', '₩2,000', '<1g'],
        ['구조', '경량 카본/발사 프레임', '1세트', '₩10,000', '₩10,000', '~15g'],
        ['가스', '헬륨 가스 (50L)', '7L 사용', '₩30,000', '₩30,000', '-'],
        ['', '', '', '1대 합계', '₩99,000', '~144g'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('■ 트래킹 시스템 (공용)')
add_table(
    ['옵션', '구성', '비용', '정밀도'],
    [
        ['Option A: OptiTrack', 'Prime 13 카메라 6대 + Motive SW', '₩20,000,000~', '±1mm'],
        ['Option B: DIY 카메라', '산업용 카메라 4대 + OpenCV + PC', '₩1,000,000~', '±2~5cm'],
    ]
)

doc.add_heading('10.3 단계별 개발 로드맵', level=2)
doc.add_paragraph('■ Phase 1: 단일 기체 프로토타입 (약 2~3개월)')
doc.add_paragraph('  • 돌고래 형상 엔벨로프 제작 및 부력 테스트')
doc.add_paragraph('  • ESP32 + IMU + 서보 4개 조립')
doc.add_paragraph('  • 자세 안정화 PID 구현 및 튜닝')
doc.add_paragraph('  • RC 조종기로 수동 비행 테스트')
doc.add_paragraph('  • Gazebo SITL 환경 구축')
doc.add_paragraph()
doc.add_paragraph('■ Phase 2: 자율 위치제어 (약 2~3개월)')
doc.add_paragraph('  • 외부 트래킹 시스템 구축 (Option B DIY 우선)')
doc.add_paragraph('  • Ground Station 위치 제어기 개발')
doc.add_paragraph('  • WiFi 통신 프로토콜 구현')
doc.add_paragraph('  • 단일 드론 자율 비행 (Waypoint 추종)')
doc.add_paragraph('  • PID 위치 제어 튜닝')
doc.add_paragraph()
doc.add_paragraph('■ Phase 3: 군집 비행 (약 2~3개월)')
doc.add_paragraph('  • 드론 5대 제작 및 개별 캘리브레이션')
doc.add_paragraph('  • 군집 제어 서버 개발')
doc.add_paragraph('  • 충돌 회피 모듈 구현')
doc.add_paragraph('  • 안무 편집기 개발 (기본 버전)')
doc.add_paragraph('  • 3대 → 5대 군집 비행 테스트')
doc.add_paragraph()
doc.add_paragraph('■ Phase 4: 아트 퍼포먼스 (약 1~2개월)')
doc.add_paragraph('  • 10대 규모 확장')
doc.add_paragraph('  • LED 조명 안무 동기화')
doc.add_paragraph('  • 음악 싱크 시스템 구현')
doc.add_paragraph('  • 리허설 및 최적화')
doc.add_paragraph('  • 첫 공연')

doc.add_page_break()

# ═══════════════════════════════════════════
# 11장. 결론 및 향후 과제
# ═══════════════════════════════════════════
doc.add_heading('11장. 결론 및 향후 과제', level=1)

doc.add_paragraph(
    '본 보고서에서는 돌고래 형상의 헬륨 부력 드론을 설계하고, '
    '실내 환경에서 5~10대 규모의 군집 아트 퍼포먼스를 구현하기 위한 '
    '기술적 분석과 설계 방향을 제시하였다.'
)
doc.add_paragraph()
doc.add_paragraph('■ 핵심 결론')
doc.add_paragraph('  1. 2m급 돌고래 형상 엔벨로프는 약 240~280g의 탑재중량을 확보할 수 있어, '
                   '경량 전자부(~30g)와 구동부(~14g) 탑재에 충분하다.')
doc.add_paragraph('  2. 꼬리지느러미 진동 추진은 약 8gf의 추진력을 발생시키며, '
                   '실내에서 0.3~0.5 m/s의 순항 속도를 달성할 수 있다.')
doc.add_paragraph('  3. 비행제어는 ESP32 기반 커스텀 FC가 적합하며, '
                   '2단 계층 PID 구조(위치+자세)로 설계한다.')
doc.add_paragraph('  4. 위치 제어는 외부 카메라 트래킹 시스템이 최적이며, '
                   '기체 무게 부담 없이 ±5cm 이내의 정밀도를 달성할 수 있다.')
doc.add_paragraph('  5. 군집 제어는 중앙 집중형 아키텍처를 채택하고, '
                   'ORCA 알고리즘으로 충돌 회피를 구현한다.')
doc.add_paragraph()
doc.add_paragraph('■ 향후 과제')
doc.add_paragraph('  • 엔벨로프 3D 패턴 설계 및 제작 공정 확립')
doc.add_paragraph('  • Gazebo 커스텀 부력/추진 플러그인 개발')
doc.add_paragraph('  • 꼬리지느러미 추진 효율의 실험적 검증 (풍동 또는 실측)')
doc.add_paragraph('  • MPC(Model Predictive Control) 기반 고급 제어기 연구')
doc.add_paragraph('  • 헬륨 가스 보충 자동화 시스템')
doc.add_paragraph('  • 관객 인터랙션 (센서 기반 반응형 안무)')
doc.add_paragraph('  • 실외 확장 시 GPS/RTK 기반 위치제어 전환')

doc.add_page_break()

# ═══════════════════════════════════════════
# 참고문헌
# ═══════════════════════════════════════════
doc.add_heading('참고문헌', level=1)

references = [
    '[1] Fish, F. E., & Lauder, G. V. (2006). "Passive and active flow control by swimming fishes and mammals." '
    'Annual Review of Fluid Mechanics, 38, 193-224.',

    '[2] Fish, F. E., & Hui, C. A. (1991). "Dolphin swimming — a review." '
    'Mammal Review, 21(4), 181-195.',

    '[3] Taylor, G. K., Nudds, R. L., & Thomas, A. L. R. (2003). "Flying and swimming animals '
    'cruise at a Strouhal number tuned for high power efficiency." Nature, 425(6959), 707-711.',

    '[4] Triantafyllou, M. S., Triantafyllou, G. S., & Yue, D. K. P. (2000). "Hydrodynamics of '
    'fishlike swimming." Annual Review of Fluid Mechanics, 32(1), 33-53.',

    '[5] Sfakiotakis, M., Lane, D. M., & Davies, J. B. C. (1999). "Review of fish swimming modes '
    'for aquatic locomotion." IEEE Journal of Oceanic Engineering, 24(2), 237-252.',

    '[6] Fish, F. E. (2004). "Structure and mechanics of nonpiscine control surfaces." '
    'IEEE Journal of Oceanic Engineering, 29(3), 605-621.',

    '[7] Lighthill, M. J. (1970). "Aquatic animal propulsion of high hydromechanical efficiency." '
    'Journal of Fluid Mechanics, 44(2), 265-301.',

    '[8] Liao, L., & Pasternak, I. (2009). "A review of airship structural research and development." '
    'Progress in Aerospace Sciences, 45(4-5), 83-96.',

    '[9] Khoury, G. A. (2012). "Airship Technology." Cambridge University Press, 2nd Edition.',

    '[10] Mueller, J., et al. (2009). "Development of an aerodynamic model and control law '
    'design for a high altitude airship." AIAA 2009-2810.',

    '[11] Colozza, A. (2003). "Initial feasibility assessment of a high altitude long endurance '
    'airship." NASA/CR-2003-212724.',

    '[12] Hoerner, S. F. (1965). "Fluid-Dynamic Drag." Published by the author.',

    '[13] Festo AG & Co. KG. (2007-2018). Bionic Learning Network Technical Reports. '
    'https://www.festo.com/group/en/cms/10156.htm',

    '[14] Preiss, J. A., et al. (2017). "Crazyswarm: A large nano-quadcopter swarm." '
    'IEEE International Conference on Robotics and Automation (ICRA).',

    '[15] Augugliaro, F., et al. (2014). "The Flight Assembled Architecture installation." '
    'IEEE Robotics & Automation Magazine, 21(3), 34-44.',

    '[16] Mellinger, D., & Kumar, V. (2011). "Minimum snap trajectory generation and control '
    'for quadrotors." IEEE ICRA 2011.',

    '[17] Van den Berg, J., et al. (2011). "Reciprocal n-body collision avoidance." '
    'Robotics Research, Springer, 3-19.',

    '[18] Ren, W., & Beard, R. W. (2008). "Distributed Consensus in Multi-vehicle Cooperative Control." '
    'Springer London.',

    '[19] Vásárhelyi, G., et al. (2018). "Optimized flocking of autonomous drones in confined environments." '
    'Science Robotics, 3(20), eaat3536.',

    '[20] Madgwick, S. O. H., Harrison, A. J. L., & Vaidyanathan, R. (2011). '
    '"Estimation of IMU and MARG orientation using a gradient descent algorithm." '
    'IEEE International Conference on Rehabilitation Robotics.',

    '[21] Åström, K. J., & Murray, R. M. (2021). "Feedback Systems: An Introduction for Scientists '
    'and Engineers." Princeton University Press, 2nd Edition.',

    '[22] Beard, R. W., & McLain, T. W. (2012). "Small Unmanned Aircraft: Theory and Practice." '
    'Princeton University Press.',

    '[23] Coulter, R. C. (1992). "Implementation of the pure pursuit path tracking algorithm." '
    'Carnegie Mellon University, CMU-RI-TR-92-01.',

    '[24] Rao, J., et al. (2003). "GTMARS: A miniature autonomous robotic airship." '
    'IEEE/RSJ International Conference on Intelligent Robots and Systems.',

    '[25] Van der Zwaan, S., et al. (2000). "Vision based station keeping and docking '
    'for an aerial blimp." IEEE/RSJ IROS 2000.',

    '[26] Fedorenko, R., & Krukhmalev, V. (2016). "Indoor autonomous airship control and navigation system." '
    'MATEC Web of Conferences, 42, 01006.',

    '[27] Krajník, T., et al. (2014). "A practical multirobot localization system." '
    'Journal of Intelligent & Robotic Systems, 76(3), 539-562.',

    '[28] Bandyopadhyay, P. R. (2005). "Trends in biorobotic autonomous undersea vehicles." '
    'IEEE Journal of Oceanic Engineering, 30(1), 109-139.',

    '[29] Shyy, W., et al. (2010). "Recent progress in flapping wing aerodynamics and aeroelasticity." '
    'Progress in Aerospace Sciences, 46(7), 284-327.',

    '[30] Li, Y., et al. (2011). "Dynamic modeling and analysis of a tri-lobed airship." '
    'Aerospace Science and Technology, 15(6), 481-488.',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(9)

# ── 저장 ──
output_path = '/Users/limchanyeong/Desktop/dolphinedrone/돌고래_헬륨_드론_기술분석_보고서.docx'
doc.save(output_path)
print(f'문서 생성 완료: {output_path}')
